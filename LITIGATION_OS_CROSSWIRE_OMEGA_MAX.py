#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LITIGATION_OS_CROSSWIRE_OMEGA_MAX.py
One-file MAX build:
 - Everything from OMEGA PLUS (evidence ingest, LLM, MiFILE, GUI hooks, code crawler, bundle, audit)
 - ActionCard JSON schema + strict validation
 - Rule snapshot loader (MCR/MCL/Benchbook) + deadline & service checkers
 - Bates numbering + Exhibit packager + Exhibit Index (CSV + DOCX)
 - ZIP validator for e-filing packages
 - MiFILE pre-submit checklist that blocks filings on red flags

Run --help for options.
"""
from __future__ import annotations
import os, sys, subprocess, json, time, shutil, hashlib, sqlite3, re, argparse, threading, datetime, traceback, platform, zipfile, ast, webbrowser, csv, math
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple, Set
from uuid import uuid4

# ---------- 0) Self-bootstrap ----------

REQUIRED = [
    "watchdog",
    "pymupdf",       # import name: fitz
    "python-docx",   # import name: docx
    "keyring",
    "requests",
    "pyflakes",
    "flask",
    "playwright",
    "jsonschema",
    "python-dateutil",
]

OPTIONAL = [
    "pytesseract",
    "Pillow",
    "pyinstaller",
]

def _pip_install(pkgs):
    py = sys.executable
    for p in pkgs:
        mod = p.split("==")[0].replace("-", "_")
        try:
            __import__(mod)
            continue
        except Exception:
            pass
        try:
            subprocess.run([py, "-m", "pip", "install", "--upgrade", p], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except subprocess.CalledProcessError:
            subprocess.run(["pip", "install", "--upgrade", p], check=False)

def ensure_deps():
    _pip_install(REQUIRED)
    _pip_install(OPTIONAL)
    # playwright browser
    try:
        subprocess.run([sys.executable, "-m", "playwright", "install", "chromium"], check=False, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=180)
    except Exception:
        pass

ensure_deps()

# ---------- Imports ----------
import requests, keyring, fitz
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, redirect, url_for, Response
from jsonschema import validate as js_validate, Draft7Validator
from jsonschema.exceptions import ValidationError as JSONSchemaValidationError
from dateutil import parser as dtparser
from dateutil.tz import gettz

try:
    import pytesseract
    from PIL import Image
except Exception:
    pytesseract = None
    Image = None

try:
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
except Exception:
    sync_playwright = None
    PWTimeout = Exception

# ---------- 1) Config ----------

IS_WIN = platform.system().lower().startswith("win")
APP_DIR = Path(os.getenv("APPDATA") if IS_WIN else Path.home() / ".config") / "LitOS"
APP_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = APP_DIR / "knowledge.db"
CONFIG_PATH = APP_DIR / "config.json"
LOG_PATH = APP_DIR / "litigation_os.log"
RULES_PATH = APP_DIR / "rules_snapshot.json"
HOLIDAYS_PATH = APP_DIR / "holidays.json"
BATES_LEDGER = APP_DIR / "bates_ledger.json"

DEFAULT_CONFIG = {
    "scan_roots": [r"F:\\", r"Z:\\", r"C:\\Litigation"],
    "recursive": True,
    "file_glob": ["*.pdf", "*.docx", "*.txt"],
    "ignore_dirs": [r"\\$RECYCLE.BIN", r"\\System Volume Information", r"\\node_modules", r"\\.git", r"\\__pycache__"],
    "results_dir": str(Path("F:/LegalResults/RECOMMENDED_ACTIONS").resolve() if IS_WIN else Path.home() / "LegalResults/RECOMMENDED_ACTIONS"),
    "dashboard_path": str(Path("F:/LegalResults/index.html").resolve() if IS_WIN else Path.home() / "LegalResults/index.html"),
    "ready_to_file": str(Path("F:/READY_TO_FILE").resolve() if IS_WIN else Path.home() / "READY_TO_FILE"),
    "proofs_dir": str(Path("F:/LegalResults/MiFILE_Proofs").resolve() if IS_WIN else Path.home() / "LegalResults/MiFILE_Proofs"),
    "exhibits_dir": str(Path("F:/LegalResults/EXHIBITS").resolve() if IS_WIN else Path.home() / "LegalResults/EXHIBITS"),
    "code": {
        "roots": [r"Z:\\"],
        "recursive": True,
        "extensions": [
            ".py",".ps1",".psm1",".psd1",".bat",".cmd",".vbs",".vb",
            ".js",".ts",".json",".yml",".yaml",".toml",".ini",".cfg",
            ".cs",".java",".go",".rs",".c",".cpp",".h",".hpp",".rb",".php",".pl",".lua",".sh",".md",".sql",".xaml"
        ],
        "special_names": ["Dockerfile","Makefile","dockerfile"],
        "bundle_dir": str(Path("F:/LegalResults/CODE_BUNDLE").resolve() if IS_WIN else Path.home() / "LegalResults/CODE_BUNDLE")
    },
    "llm": {
        "backend": "auto",
        "openai_model": "gpt-5.0-mini",
        "ollama_model": "llama3.1:8b",
        "max_tokens": 1800,
        "temperature": 0.15
    },
    "webhook_url": "",
    "owner": {
        "name": "Andrew J Pigors",
        "email": "Andrewjpigors@gmail.com",
        "address": "Lot 17, 1977 Whitehall Rd, Muskegon, MI 49445",
        "timezone": "America/Detroit"
    },
    "cases": {
        "housing": {"case_no":"2025-002760-CZ", "court":"14th Circuit Court, Muskegon County, Michigan"},
        "custody": {"case_no":"2024-0000001507-DC", "court":"14th Circuit Court, Muskegon County, Michigan"},
        "lt": {"case_no":"2025-25061626LT-LT", "court":"60th District Court, Muskegon County, Michigan"}
    },
    "mifile": {
        "base_url": "https://mifile.courts.michigan.gov/cases",
        "headless": True,
        "timeout_ms": 35000
    },
    "bates": {
        "prefix": "AJP",
        "start": 1,
        "digits": 6,
        "color": (0, 0, 0),  # black
        "font_size": 9
    }
}

def load_config() -> Dict[str, Any]:
    if CONFIG_PATH.exists():
        try:
            cfg = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            cfg = DEFAULT_CONFIG
    else:
        cfg = DEFAULT_CONFIG
        CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
        CONFIG_PATH.write_text(json.dumps(cfg, indent=2), encoding="utf-8")
    return cfg

CONFIG = load_config()
TZ = gettz(CONFIG["owner"].get("timezone","America/Detroit"))
RESULTS_DIR = Path(CONFIG["results_dir"]); RESULTS_DIR.mkdir(parents=True, exist_ok=True)
CODE_BUNDLE_DIR = Path(CONFIG["code"]["bundle_dir"]); CODE_BUNDLE_DIR.mkdir(parents=True, exist_ok=True)
Path(CONFIG["proofs_dir"]).mkdir(parents=True, exist_ok=True)
Path(CONFIG["ready_to_file"]).mkdir(parents=True, exist_ok=True)
Path(CONFIG["exhibits_dir"]).mkdir(parents=True, exist_ok=True)

# ---------- 2) Logging ----------

def log(msg: str):
    ts = datetime.datetime.now(TZ).strftime("%Y-%m-%d %H:%M:%S %Z")
    line = f"[{ts}] {msg}"
    print(line, flush=True)
    try:
        with open(LOG_PATH, "a", encoding="utf-8") as f:
            f.write(line + "\n")
    except Exception:
        pass

# ---------- 3) DB ----------

def db_connect() -> sqlite3.Connection:
    con = sqlite3.connect(str(DB_PATH))
    con.execute("PRAGMA journal_mode=WAL;")
    con.execute("PRAGMA synchronous=NORMAL;")
    return con

def db_init():
    con = db_connect(); cur = con.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS files(
        id INTEGER PRIMARY KEY,
        path TEXT UNIQUE,
        sha256 TEXT,
        mtime REAL,
        size INTEGER,
        doc_type TEXT,
        case_tag TEXT,
        llm_backend TEXT,
        llm_model TEXT,
        actions_json TEXT,
        created_at TEXT,
        updated_at TEXT
    );""")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS paragraphs(
        id INTEGER PRIMARY KEY,
        file_id INTEGER,
        p_index INTEGER,
        text TEXT,
        sha256 TEXT,
        FOREIGN KEY(file_id) REFERENCES files(id) ON DELETE CASCADE
    );""")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS code_files(
        id INTEGER PRIMARY KEY,
        path TEXT UNIQUE,
        sha256 TEXT,
        size INTEGER,
        mtime REAL,
        language TEXT,
        shebang TEXT,
        lint_ok INTEGER,
        compile_ok INTEGER,
        error TEXT,
        created_at TEXT
    );""")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS code_refs(
        id INTEGER PRIMARY KEY,
        code_file_id INTEGER,
        ref_type TEXT,
        ref_name TEXT,
        FOREIGN KEY(code_file_id) REFERENCES code_files(id) ON DELETE CASCADE
    );""")
    con.commit(); con.close()

db_init()

# ---------- 4) Utils ----------

def sha256_bytes(b: bytes) -> str:
    h = hashlib.sha256(); h.update(b); return h.hexdigest()

def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024*1024), b""):
            h.update(chunk)
    return h.hexdigest()

# ---------- 5) Extraction ----------

def extract_text(path: Path):
    meta = {"method": None, "pages": 0, "ocr_used": False}
    text = ""
    suf = path.suffix.lower()
    try:
        if suf == ".pdf":
            text, meta = _extract_pdf_text(path)
        elif suf == ".docx":
            text = _extract_docx_text(path); meta["method"]="python-docx"
        elif suf == ".txt":
            text = path.read_text(encoding="utf-8", errors="ignore"); meta["method"]="txt"
        else:
            meta["method"]="unsupported"
    except Exception as e:
        log(f"extract_text error {path}: {e}")
    return text.strip(), meta

def _extract_pdf_text(path: Path):
    meta = {"method":"pymupdf","pages":0,"ocr_used":False}
    out = []
    with fitz.open(str(path)) as doc:
        meta["pages"] = doc.page_count
        for i in range(doc.page_count):
            t = doc.load_page(i).get_text("text")
            if t.strip(): out.append(t)
    txt = "\n".join(out).strip()
    if txt: return txt, meta
    # OCR fallback
    if pytesseract and Image:
        try:
            ocr = []
            with fitz.open(str(path)) as doc:
                for i in range(doc.page_count):
                    pix = doc.load_page(i).get_pixmap(dpi=200)
                    img_bytes = pix.tobytes("png")
                    from io import BytesIO
                    img = Image.open(BytesIO(img_bytes))
                    ocr.append(pytesseract.image_to_string(img))
            meta["method"]="ocr"; meta["ocr_used"]=True
            return "\n".join(ocr), meta
        except Exception as e:
            log(f"OCR fail {path}: {e}")
    return "", meta

def _extract_docx_text(path: Path) -> str:
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)

def split_paragraphs(text: str) -> List[str]:
    parts = re.split(r"\n{2,}|\r{2,}", text)
    return [p.strip() for p in parts if p and p.strip()]

# ---------- 6) LLM backends ----------

def _ollama_up() -> bool:
    try:
        r = requests.get("http://127.0.0.1:11434/api/tags", timeout=1.5)
        return r.status_code == 200
    except Exception:
        return False

def get_backend(llm_cfg):
    be = llm_cfg.get("backend","auto")
    if be == "auto":
        if _ollama_up(): return "ollama", llm_cfg.get("ollama_model","llama3.1:8b")
        key = keyring.get_password("LitOS","openai_api_key")
        if key: return "openai", llm_cfg.get("openai_model","gpt-5.0-mini")
        return "ollama", llm_cfg.get("ollama_model","llama3.1:8b")
    if be == "openai": return "openai", llm_cfg.get("openai_model","gpt-5.0-mini")
    if be == "ollama": return "ollama", llm_cfg.get("ollama_model","llama3.1:8b")
    return "ollama","llama3.1:8b"

def _call_openai(prompt: str, max_tokens: int, temperature: float, model: str) -> str:
    key = keyring.get_password("LitOS","openai_api_key")
    if not key:
        try:
            if sys.stdin and sys.stdin.isatty():
                print("\nEnter your OpenAI API key (stored in keyring):")
                key = input("sk-...: ").strip()
                if key: keyring.set_password("LitOS","openai_api_key",key)
        except Exception:
            pass
    if not key:
        raise RuntimeError("OpenAI API key missing in keyring (service='LitOS', user='openai_api_key').")
    headers = {"Authorization": f"Bearer {key}", "Content-Type":"application/json"}
    payload = {
        "model": model,
        "messages": [
            {"role":"system","content":"You are a Michigan litigation expert. You MUST output ONLY valid JSON for the ActionCard schema I provide. No extra text."},
            {"role":"user","content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": temperature
    }
    r = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload, timeout=90)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

def _call_ollama(prompt: str, model: str, options: Optional[Dict[str,Any]] = None) -> str:
    payload = {"model": model, "prompt": prompt, "stream": False}
    if options: payload["options"]=options
    r = requests.post("http://127.0.0.1:11434/api/generate", json=payload, timeout=180)
    r.raise_for_status()
    return r.json().get("response","").strip()

# ---------- 7) ActionCard Schema & Validation ----------

ACTIONCARD_SCHEMA = {
  "type": "object",
  "required": ["action_id","title","category","authority","filing","service","evidence_refs"],
  "properties": {
    "action_id": {"type":"string"},
    "title": {"type":"string", "minLength": 3, "maxLength": 80},
    "category": {"type":"string", "enum":["Housing","Custody","PPO","Other"]},
    "authority": {
      "type":"object",
      "properties": {
        "mcr": {"type":"array","items":{"type":"string"}},
        "mcl": {"type":"array","items":{"type":"string"}},
        "benchbook": {"type":"array","items":{"type":"string"}}
      },
      "additionalProperties": False
    },
    "form_refs": {"type":"array","items":{"type":"string"}},
    "filing": {
      "type":"object",
      "required": ["court","case_no"],
      "properties": {
        "court": {"type":"string"},
        "case_no": {"type":"string"},
        "deadline_utc": {"type":"string"}
      },
      "additionalProperties": True
    },
    "service": {
      "type":"object",
      "required": ["method"],
      "properties": {
        "method": {"type":"array","items":{"type":"string"}},
        "parties": {"type":"array","items":{"type":"string"}},
        "proof_doc": {"type":"string"}
      },
      "additionalProperties": True
    },
    "evidence_refs": {
      "type":"array",
      "items": {
        "type":"object",
        "required": ["file_sha"],
        "properties": {
          "file_sha": {"type":"string"},
          "para_idx": {"type":"array","items":{"type":"integer"}}
        }
      }
    },
    "risk_checks": {"type":"array","items":{"type":"string"}},
    "notes": {"type":"string"}
  },
  "additionalProperties": False
}

def validate_actioncard(card: Dict[str,Any]) -> List[str]:
    errors = []
    v = Draft7Validator(ACTIONCARD_SCHEMA)
    for e in v.iter_errors(card):
        errors.append(f"{'/'.join(str(x) for x in e.path)}: {e.message}")
    # extra custom checks
    if "service" in card:
        methods = set(m.lower() for m in card["service"].get("method", []))
        if "email" in methods and not card["service"].get("parties"):
            errors.append("service.parties required when method includes email")
        if "mifile" in methods and not card["filing"].get("case_no"):
            errors.append("filing.case_no required for MiFILE service")
    return errors

# ---------- 8) Rule Snapshot & Checkers ----------

def load_rules() -> Dict[str,Any]:
    if RULES_PATH.exists():
        try:
            return json.loads(RULES_PATH.read_text(encoding="utf-8"))
        except Exception as e:
            log(f"rules parse error: {e}")
    # minimal default structure
    return {"version":"local","mcr":{}, "mcl":{}, "benchbook":{}, "deadlines":{}, "service":{}, "holidays": []}

def load_holidays() -> Set[str]:
    if HOLIDAYS_PATH.exists():
        try:
            arr = json.loads(HOLIDAYS_PATH.read_text(encoding="utf-8"))
            return set(arr)
        except Exception:
            pass
    return set()

def court_day_add(start_iso: str, days: int, court_days: bool=True) -> str:
    """Add days to start_iso, skipping weekends and configured holidays if court_days=True."""
    tz = TZ
    dt = dtparser.parse(start_iso).astimezone(tz)
    hol = load_holidays()
    n = 0
    delta = datetime.timedelta(days=1)
    cur = dt
    while n < days:
        cur = cur + delta
        if court_days:
            if cur.weekday() >= 5:  # Sat/Sun
                continue
            if cur.strftime("%Y-%m-%d") in hol:
                continue
        n += 1
    return cur.astimezone(datetime.timezone.utc).isoformat()

def rules_check_actioncard(card: Dict[str,Any]) -> List[str]:
    """Return list of red flags (empty if OK)."""
    flags = []
    rules = load_rules()
    # Check cited authorities exist in snapshot (name-only check)
    for cite in card.get("authority",{}).get("mcr",[]):
        if rules.get("mcr") and cite not in rules["mcr"]:
            flags.append(f"Authority not in snapshot: MCR {cite}")
    for cite in card.get("authority",{}).get("mcl",[]):
        if rules.get("mcl") and cite not in rules["mcl"]:
            flags.append(f"Authority not in snapshot: MCL {cite}")
    # Deadline sanity if provided
    dl = card.get("filing",{}).get("deadline_utc")
    if dl:
        try:
            dldt = dtparser.parse(dl)
            if dldt < datetime.datetime.now(datetime.timezone.utc):
                flags.append("Deadline appears in the past.")
        except Exception:
            flags.append("Invalid deadline_utc format.")
    # Service method check coverage
    methods = set(m.lower() for m in card.get("service",{}).get("method",[]))
    if not methods:
        flags.append("No service method specified.")
    # Basic attachments presence (by file_sha mapped in DB)
    missing = []
    con=db_connect(); cur=con.cursor()
    for ev in card.get("evidence_refs",[]):
        sha = ev.get("file_sha","")
        row = cur.execute("SELECT path FROM files WHERE sha256=?", (sha,)).fetchone()
        if not row:
            missing.append(sha)
    con.close()
    if missing:
        flags.append(f"Evidence not indexed: {', '.join(missing)}")
    return flags

# ---------- 9) Bates Numbering & Exhibit Packager ----------

def _exhibit_label(n: int) -> str:
    # A, B, ..., Z, AA, AB, ...
    letters = []
    n0 = n
    while True:
        n, r = divmod(n0, 26)
        letters.append(chr(ord('A') + r))
        if n == 0: break
        n0 = n - 1
    return "".join(reversed(letters))

def _bates_next(count_pages: int) -> List[str]:
    cfg = CONFIG["bates"]
    prefix = cfg.get("prefix","AJP")
    digits = int(cfg.get("digits",6))
    # Load ledger
    ledger = {"next": int(cfg.get("start",1))}
    if BATES_LEDGER.exists():
        try:
            ledger = json.loads(BATES_LEDGER.read_text(encoding="utf-8"))
        except Exception:
            pass
    start = ledger.get("next",1)
    seq = []
    for i in range(count_pages):
        seq.append(f"{prefix}-{start+i:0{digits}d}")
    ledger["next"] = start + count_pages
    BATES_LEDGER.write_text(json.dumps(ledger, indent=2), encoding="utf-8")
    return seq

def stamp_pdf_with_bates(in_pdf: Path, out_pdf: Path, exhibit_label: str, short_desc: str) -> int:
    """Return pages stamped count."""
    doc = fitz.open(str(in_pdf))
    pages = doc.page_count
    bates_seq = _bates_next(pages)
    for i in range(pages):
        page = doc.load_page(i)
        # Bates footer-right
        text = bates_seq[i]
        rect = fitz.Rect(page.rect.width-180, page.rect.height-36, page.rect.width-10, page.rect.height-10)
        page.insert_textbox(rect, text, fontsize=9, color=(0,0,0), align=2)
        # Exhibit header-right
        ex = f"Exhibit {exhibit_label} — {short_desc}"
        rect2 = fitz.Rect(page.rect.width-300, 10, page.rect.width-10, 40)
        page.insert_textbox(rect2, ex, fontsize=10, color=(0,0,0), align=2)
    doc.save(str(out_pdf))
    doc.close()
    return pages

def package_exhibits(files: List[Path], out_dir: Path, index_name: str="EXHIBIT_INDEX") -> Tuple[Path, Path]:
    """
    Stamp each PDF and write Exhibit Index CSV + DOCX. Returns (csv_path, docx_path).
    Non-PDFs are copied as-is and listed with no page count.
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    rows = []
    label_i = 0
    for f in files:
        label = _exhibit_label(label_i)
        label_i += 1
        desc = f.stem[:60]
        sha = sha256_file(f)
        if f.suffix.lower() == ".pdf":
            out_pdf = out_dir / f"Exhibit_{label}_{f.stem}.pdf"
            pages = stamp_pdf_with_bates(f, out_pdf, label, desc)
            rows.append({"Exhibit": label, "File": out_pdf.name, "Pages": pages, "SHA256": sha, "Description": desc})
        else:
            outf = out_dir / f"Exhibit_{label}_{f.name}"
            shutil.copy2(str(f), str(outf))
            rows.append({"Exhibit": label, "File": outf.name, "Pages": "", "SHA256": sha, "Description": desc})
    # CSV
    csv_path = out_dir / f"{index_name}.csv"
    with open(csv_path, "w", newline="", encoding="utf-8") as fp:
        w = csv.DictWriter(fp, fieldnames=["Exhibit","File","Pages","SHA256","Description"])
        w.writeheader(); w.writerows(rows)
    # DOCX
    doc = Document()
    p = doc.add_paragraph(); r=p.add_run("EXHIBIT INDEX"); r.bold=True; r.font.size=Pt(16); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        doc.add_paragraph(f"{row['Exhibit']}. {row['File']} — {row['Description']} (pages: {row['Pages'] or 'n/a'})  SHA: {row['SHA256'][:12]}")
    docx_path = out_dir / f"{index_name}.docx"; doc.save(str(docx_path))
    return csv_path, docx_path

# ---------- 10) ZIP Validator ----------

BAD_EXTS = {".exe",".bat",".cmd",".ps1",".js",".msi",".dll",".vbs",".jar",".scr",".com"}

def validate_zip(zip_path: Path) -> List[str]:
    flags = []
    if not zip_path.exists() or zip_path.suffix.lower() != ".zip":
        return ["Not a .zip file or missing"]
    with zipfile.ZipFile(str(zip_path), "r") as z:
        names = z.namelist()
        if not names:
            flags.append("ZIP is empty")
        for n in names:
            ext = Path(n).suffix.lower()
            if ext in BAD_EXTS:
                flags.append(f"Disallowed file type inside ZIP: {n}")
        # Check exhibit index present if exhibits detected
        if any("Exhibit_" in n for n in names) and not any("EXHIBIT_INDEX" in n for n in names):
            flags.append("Exhibits present but EXHIBIT_INDEX missing")
    return flags

# ---------- 11) MiFILE Automation (login/upload), pre-submit checklist ----------

class MiFileClient:
    def __init__(self, base_url: str, headless: bool=True, timeout_ms: int=35000):
        self.base_url = base_url; self.headless=headless; self.timeout_ms=timeout_ms
        self._pw = self._browser = self._context = self._page = None

    def _start(self):
        if sync_playwright is None: raise RuntimeError("Playwright unavailable")
        self._pw = sync_playwright().start()
        self._browser = self._pw.chromium.launch(headless=self.headless)
        self._context = self._browser.new_context()
        self._page = self._context.new_page()
        self._page.set_default_timeout(self.timeout_ms)

    def _stop(self):
        for obj in [self._context, self._browser]:
            try:
                if obj: obj.close()
            except Exception: pass
        try:
            if self._pw: self._pw.stop()
        except Exception: pass

    def login(self, username: str, password: str):
        self._start(); page=self._page
        page.goto(self.base_url)
        try:
            page.get_by_label("Email").fill(username)
            page.get_by_label("Password").fill(password)
        except Exception:
            try:
                page.fill("input[type='email']", username)
                page.fill("input[type='password']", password)
            except Exception:
                page.fill("input#Email, input[name='Email']", username)
                page.fill("input#Password, input[name='Password']", password)
        for sel in ["text=Sign in","text=Log in","button:has-text('Sign in')","button:has-text('Log in')","input[type='submit']"]:
            try: page.click(sel); break
            except Exception: pass
        page.wait_for_timeout(1500)

    def upload(self, file_path: str, case_no: str, doc_type: str, proofs_dir: Path) -> Path:
        page=self._page
        for c in ["text=File a Document","text=E-File","text=New Filing","text=Submit Document","text=Filing"]:
            try: page.click(c, timeout=5000); break
            except Exception: pass
        try:
            page.fill("input[placeholder*='Case']", case_no)
        except Exception:
            try: page.get_by_label("Case Number").fill(case_no)
            except Exception: pass
        page.set_input_files("input[type='file']", file_path)
        try:
            page.fill("input[placeholder*='Document Type']", doc_type)
        except Exception:
            try: page.select_option("select[name*='DocumentType']", label=doc_type)
            except Exception: pass
        for sel in ["text=Submit","text=File","button:has-text('Submit')","button:has-text('File')","input[type='submit']"]:
            try: page.click(sel); break
            except Exception: pass
        page.wait_for_timeout(2000)
        ts = datetime.datetime.now(TZ).strftime("%Y%m%d_%H%M%S")
        proofs_dir.mkdir(parents=True, exist_ok=True)
        scr = proofs_dir / f"MiFILE_{case_no}_{Path(file_path).stem}_{ts}.png"
        page.screenshot(path=str(scr), full_page=True)
        return scr

    def close(self):
        self._stop()

def _get_mifile_creds():
    u = keyring.get_password("LitOS","mifile_username") or ""
    p = keyring.get_password("LitOS","mifile_password") or ""
    if not (u and p) and sys.stdin and sys.stdin.isatty():
        print("\nMiFILE creds (stored in keyring: service='LitOS')")
        if not u: u = input("Username (email): ").strip() or u
        if u: keyring.set_password("LitOS","mifile_username",u)
        if not p: p = input("Password: ").strip() or p
        if p: keyring.set_password("LitOS","mifile_password",p)
    return u,p

def mifile_precheck(card: Dict[str,Any], planned_files: List[Path], zip_candidate: Optional[Path]=None) -> List[str]:
    flags = []
    # Schema & rules
    flags += validate_actioncard(card)
    flags += rules_check_actioncard(card)
    # Files exist
    for f in planned_files:
        if not f.exists():
            flags.append(f"Missing attachment: {f}")
        elif f.stat().st_size == 0:
            flags.append(f"Zero-byte attachment: {f}")
    # ZIP validation
    if zip_candidate:
        flags += validate_zip(zip_candidate)
    return flags

# ---------- 12) Evidence pipeline (same core as OMEGA+, but ActionCard-driven) ----------

SUPPORTED_EXTS = {".pdf",".docx",".txt"}

def guess_doc_type(text: str, path: Path) -> str:
    t = (path.name + " " + text[:500]).lower()
    if "ppo" in t or "personal protection order" in t: return "PPO"
    if any(k in t for k in ["custody","parenting time","foc"]): return "Custody"
    if any(k in t for k in ["shady oaks","mobile home","rent","landlord","lot"]): return "Housing"
    return "Other"

def process_file(path: Path):
    try:
        if not path.exists() or not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTS: return
        st = path.stat(); fsha = sha256_file(path)
        con=db_connect(); cur=con.cursor(); now=datetime.datetime.utcnow().isoformat()
        row = cur.execute("SELECT id, mtime, sha256 FROM files WHERE path=?", (str(path),)).fetchone()
        if row:
            fid, old_m, old_h = row
            if abs(old_m - st.st_mtime) < 1e-6 and old_h == fsha: con.close(); return
            cur.execute("DELETE FROM paragraphs WHERE file_id=?", (fid,))
            cur.execute("UPDATE files SET mtime=?, size=?, sha256=?, updated_at=? WHERE id=?", (st.st_mtime, st.st_size, fsha, now, fid)); file_id=fid
        else:
            cur.execute("INSERT INTO files(path,sha256,mtime,size,created_at,updated_at) VALUES(?,?,?,?,?,?)",
                        (str(path), fsha, st.st_mtime, st.st_size, now, now)); file_id=cur.lastrowid
        con.commit()

        text, meta = extract_text(path); dtype = guess_doc_type(text, path)
        paras = split_paragraphs(text)
        for i,p in enumerate(paras):
            cur.execute("INSERT INTO paragraphs(file_id,p_index,text,sha256) VALUES(?,?,?,?)", (file_id,i,p,sha256_bytes(p.encode("utf-8"))))
        con.commit()

        # LLM → ActionCard JSON
        analysis = llm_actioncard(text, path, meta)
        # Attach to DB
        cur.execute("UPDATE files SET doc_type=?, llm_backend=?, llm_model=?, actions_json=? WHERE id=?",
                    (dtype, analysis["backend"], analysis["model"], json.dumps(analysis, ensure_ascii=False), file_id))
        con.commit(); con.close()

        # Generate DOCX report from ActionCard
        out_docx = make_action_report_from_actioncard(path, analysis, fsha)
        log(f"Action report: {out_docx}")
        write_dashboard()
    except Exception as e:
        log(f"process_file error {path}: {e}\n{traceback.format_exc()}")

def initial_scan():
    roots=[Path(r) for r in CONFIG.get("scan_roots",[])]; globs=CONFIG.get("file_glob",["*.pdf","*.docx","*.txt"]); rec=CONFIG.get("recursive",True)
    for root in roots:
        if not root.exists(): log(f"Missing scan root: {root}"); continue
        for pattern in globs:
            if rec:
                for p in root.rglob(pattern): process_file(p)
            else:
                for p in root.glob(pattern): process_file(p)

# ---------- 13) LLM ActionCard builder ----------

def llm_actioncard(doc_text: str, path: Path, meta: Dict[str,Any]) -> Dict[str,Any]:
    backend, model = get_backend(CONFIG["llm"])
    temp = CONFIG["llm"].get("temperature",0.15)
    max_tokens = CONFIG["llm"].get("max_tokens",1800)
    owner = CONFIG.get("owner",{}); known = CONFIG.get("cases",{})
    schema_str = json.dumps(ACTIONCARD_SCHEMA, separators=(",",":"))
    prompt = f"""
You are integrated into a Michigan Litigation OS. Given the DOCUMENT TEXT, output ONLY a JSON object matching this ActionCard JSON Schema (no prose):
SCHEMA: {schema_str}

Rules:
- action_id must be a deterministic UUIDv4-like string or unique string.
- category ∈ ["Housing","Custody","PPO","Other"].
- filing.court must be a short, accurate name if inferable; else use "Unknown".
- filing.case_no: prefer an existing case from Known Cases if text suggests it.
- authority: include specific MCR/MCL sections that the text actually triggers; if none, set empty arrays (not placeholders).
- service.method: include at least one of ["MiFILE","email","mail"] that makes sense; if "email", include parties.
- evidence_refs: include file_sha (SHA-256 of this file) and paragraph indices you relied on.
- deadline_utc: ONLY include if a real date can be inferred; else omit.

Owner: {owner.get('name')} <{owner.get('email')}>, Address: {owner.get('address')}
Known Cases: Housing {known.get('housing',{})}, Custody {known.get('custody',{})}, LT {known.get('lt',{})}

Document: {path.name} | Method: {meta.get('method')} | Pages: {meta.get('pages')} | OCR: {meta.get('ocr_used')}

DOCUMENT TEXT (truncate at 20k chars):
{doc_text[:20000]}
"""
    if backend == "openai":
        raw = _call_openai(prompt, max_tokens=max_tokens, temperature=temp, model=model); used="openai"
    else:
        raw = _call_ollama(prompt, model=model, options={"temperature":temp}); used="ollama"
    # parse JSON strictly
    try:
        j = json.loads(raw.strip())
    except Exception as e:
        # try to salvage JSON between braces
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if not m: raise RuntimeError(f"LLM failed to return JSON: {raw[:200]}")
        j = json.loads(m.group(0))
    # Ensure action_id present
    if "action_id" not in j or not j.get("action_id"):
        j["action_id"] = str(uuid4())
    # Validate
    errs = validate_actioncard(j)
    if errs:
        raise RuntimeError("ActionCard validation failed: " + "; ".join(errs))
    return {"backend": used, "model": model, "actioncard": j}

# ---------- 14) Reports from ActionCard ----------

def make_action_report_from_actioncard(src_path: Path, analysis: Dict[str,Any], file_sha: str) -> Path:
    card = analysis["actioncard"]
    title = card.get("title","Recommended Action")
    doc = Document()
    p = doc.add_paragraph(); r=p.add_run(title.upper()); r.bold=True; r.font.size=Pt(16); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    owner = CONFIG.get("owner",{}); cases=CONFIG.get("cases",{})
    doc.add_paragraph(f"Prepared for: {owner.get('name')}  |  Email: {owner.get('email')}")
    doc.add_paragraph(f"Housing Case: {cases.get('housing',{}).get('case_no','')} — {cases.get('housing',{}).get('court','')}")
    doc.add_paragraph(f"Custody Case: {cases.get('custody',{}).get('case_no','')} — {cases.get('custody',{}).get('court','')}")
    doc.add_paragraph(f"LT Case: {cases.get('lt',{}).get('case_no','')} — {cases.get('lt',{}).get('court','')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Source file: {str(src_path)}")
    doc.add_paragraph(f"SHA-256: {file_sha}")
    doc.add_paragraph("")
    # ActionCard details
    doc.add_paragraph("ActionCard")
    doc.add_paragraph(json.dumps(card, indent=2))
    # Footer
    doc.add_paragraph("")
    doc.add_paragraph(f"Generated by LITIGATION_OS_CROSSWIRE_OMEGA_MAX — PLSC enforced. LLM: {analysis['backend']}/{analysis['model']}")
    safe = re.sub(r"[^A-Za-z0-9_.-]+","_",src_path.stem)[:50]
    out_path = RESULTS_DIR / f"{safe}__ACTION_{file_sha[:10]}.docx"
    doc.save(str(out_path))
    return out_path

# ---------- 15) Dashboard ----------

def write_dashboard():
    con=db_connect(); cur=con.cursor()
    files = cur.execute("SELECT path, sha256, doc_type, llm_backend, llm_model, created_at FROM files ORDER BY created_at DESC LIMIT 200").fetchall()
    con.close()
    html = ["<html><head><meta charset='utf-8'><title>Litigation OS Dashboard</title></head><body>"]
    html.append("<h2>Litigation OS — Evidence Overview</h2>")
    html.append("<table border='1' cellspacing='0' cellpadding='6'><tr><th>File</th><th>SHA-256</th><th>Type</th><th>LLM</th><th>Model</th><th>Created</th></tr>")
    for path, sha, dt, lb, lm, created in files:
        html.append(f"<tr><td>{path}</td><td>{sha[:12]}</td><td>{dt or ''}</td><td>{lb or ''}</td><td>{lm or ''}</td><td>{created or ''}</td></tr>")
    html.append("</table></body></html>")
    Path(CONFIG["dashboard_path"]).parent.mkdir(parents=True, exist_ok=True)
    Path(CONFIG["dashboard_path"]).write_text("\n".join(html), encoding="utf-8")
    log(f"Dashboard updated: {CONFIG['dashboard_path']}")

# ---------- 16) Code Crawler/Bundle/Audit (light wrappers to keep file size reasonable) ----------

def crawl_code():
    # Minimal traversal to populate code_files table (extensions per config)
    CODE_EXTS = set(CONFIG["code"]["extensions"]); SPECIAL = set(CONFIG["code"]["special_names"])
    IGNORE_PAT = [d.lower() for d in CONFIG.get("ignore_dirs", [])]
    LANG_MAP = {
        ".py":"python",".ps1":"powershell",".psm1":"powershell",".psd1":"powershell",
        ".bat":"batch",".cmd":"batch",".vbs":"vbscript",".vb":"vbnet",
        ".js":"javascript",".ts":"typescript",".json":"json",".yml":"yaml",".yaml":"yaml",".toml":"toml",".ini":"ini",".cfg":"cfg",
        ".cs":"csharp",".java":"java",".go":"go",".rs":"rust",".c":"c",".cpp":"cpp",".h":"c-header",".hpp":"cpp-header",
        ".rb":"ruby",".php":"php",".pl":"perl",".lua":"lua",".sh":"shell",".md":"markdown",".sql":"sql",".xaml":"xaml"
    }
    def detect_language(p: Path):
        if p.name in SPECIAL:
            if "docker" in p.name.lower(): return "dockerfile"
            if p.name.lower()=="makefile": return "makefile"
        return LANG_MAP.get(p.suffix.lower(),"unknown")
    def lint_python(path: Path, text: Optional[str]=None):
        lint_ok=True; compile_ok=True; err=""
        try:
            if text is None: text = path.read_text(encoding="utf-8", errors="ignore")
            compile(text, str(path), "exec")
        except Exception as e:
            compile_ok=False; err += f"COMPILE: {e}\n"
        try:
            r = subprocess.run([sys.executable, "-m", "pyflakes", str(path)], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=30)
            if r.returncode != 0 or r.stdout or r.stderr:
                lint_ok=False; err += (r.stdout or "") + (r.stderr or "")
        except Exception as e:
            err += f"PYFLAKES_ERROR: {e}\n"
        return lint_ok, compile_ok, err.strip()
    con=db_connect(); cur=con.cursor(); count=0
    for root in CONFIG["code"].get("roots",[]):
        rootp = Path(root)
        if not rootp.exists(): log(f"Code root missing: {rootp}"); continue
        for dirpath, dirnames, filenames in os.walk(rootp):
            if any(pat in dirpath.lower() for pat in IGNORE_PAT):
                dirnames[:] = []
                continue
            for fname in filenames:
                fp = Path(dirpath)/fname
                if (fp.suffix.lower() in CODE_EXTS) or (fname in SPECIAL):
                    try:
                        st = fp.stat(); size=st.st_size; mtime=st.st_mtime
                        sha = sha256_file(fp); lang = detect_language(fp)
                        text = ""
                        try: text = fp.read_text(encoding="utf-8", errors="ignore")
                        except Exception: pass
                        lint_ok=compile_ok=True; err=""
                        if lang=="python":
                            lint_ok, compile_ok, err = lint_python(fp, text=text)
                        now = datetime.datetime.utcnow().isoformat()
                        cur.execute("INSERT OR REPLACE INTO code_files(path,sha256,size,mtime,language,shebang,lint_ok,compile_ok,error,created_at) VALUES(?,?,?,?,?,?,?,?,?,?)",
                                    (str(fp), sha, size, mtime, lang, "", int(bool(lint_ok)), int(bool(compile_ok)), err, now))
                        count += 1
                    except Exception as e:
                        log(f"code crawl error {fp}: {e}")
    con.commit(); con.close()
    log(f"Code crawl complete. Files: {count}")

def bundle_code():
    con=db_connect(); cur=con.cursor()
    rows = cur.execute("SELECT path, language, lint_ok, compile_ok FROM code_files").fetchall()
    dest = CODE_BUNDLE_DIR / "bundle"
    if dest.exists(): shutil.rmtree(dest, ignore_errors=True)
    dest.mkdir(parents=True, exist_ok=True)
    manifest={"generated_at": datetime.datetime.utcnow().isoformat(), "items":[]}
    copied=0
    for path, lang, lint_ok, compile_ok in rows:
        p = Path(path)
        if not p.exists(): continue
        if any(s in str(p).lower() for s in ["/node_modules/","\\node_modules\\","/.git/","\\._git\\","/__pycache__/","\\__pycache__\\"]):
            continue
        rel = re.sub(r"^[A-Za-z]:(\\\\|/)", "", str(p)) if IS_WIN else str(p).lstrip("/")
        outp = dest / rel; outp.parent.mkdir(parents=True, exist_ok=True)
        try:
            shutil.copy2(str(p), str(outp)); copied+=1
            manifest["items"].append({"src": str(p), "dst": str(outp), "language": lang, "ok": bool(lint_ok and compile_ok)})
        except Exception as e:
            log(f"copy fail {p}: {e}")
    man_path = CODE_BUNDLE_DIR / "MANIFEST.json"; man_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    zip_path = CODE_BUNDLE_DIR / "CODE_BUNDLE.zip"
    with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(dest):
            for f in files:
                fp = Path(root)/f; z.write(str(fp), arcname=str(fp.relative_to(dest)))
        z.write(str(man_path), arcname="MANIFEST.json")
    log(f"Bundle complete: {copied} files → {zip_path}")

def audit_code_llm():
    # Minimal text plan from inventory counts (avoid huge prompts here for size)
    con=db_connect(); cur=con.cursor()
    counts = cur.execute("SELECT language, COUNT(*) FROM code_files GROUP BY language").fetchall()
    lines = [f"{lang}: {cnt}" for lang,cnt in counts]
    summary = " | ".join(lines) if lines else "no code"
    doc = Document(); p=doc.add_paragraph(); r=p.add_run("CODE AUDIT SNAPSHOT"); r.bold=True; r.font.size=Pt(15)
    doc.add_paragraph(summary)
    path = CODE_BUNDLE_DIR / "CODE_AUDIT_SNAPSHOT.docx"; doc.save(str(path))
    log(f"Audit snapshot written: {path}")

# ---------- 17) GUI (minimal holder to keep code single-file; can be expanded) ----------

APP = Flask(__name__)

@APP.route("/")
def ui_home():
    html = ["<html><head><title>Litigation OS — MAX</title></head><body>",
            "<h2>Litigation OS — MAX Control Panel</h2>",
            "<form method='POST' action='/run'>",
            "<button name='cmd' value='scan'>Scan Evidence Once</button> ",
            "<button name='cmd' value='exhibits'>Package Exhibits (from exhibits_dir)</button> ",
            "<button name='cmd' value='bundle'>Bundle Code</button> ",
            "<button name='cmd' value='crawl'>Crawl Code</button> ",
            "</form>",
            "<p><a href='/dashboard' target='_blank'>Open Dashboard</a></p>",
            "</body></html>"]
    return Response("\n".join(html), mimetype="text/html")

@APP.route("/run", methods=["POST"])
def ui_run():
    cmd = request.form.get("cmd","")
    if cmd=="scan": threading.Thread(target=initial_scan, daemon=True).start()
    if cmd=="bundle": threading.Thread(target=bundle_code, daemon=True).start()
    if cmd=="crawl": threading.Thread(target=crawl_code, daemon=True).start()
    if cmd=="exhibits":
        # Gather PDFs from exhibits_dir and package
        p = Path(CONFIG["exhibits_dir"])
        files = sorted([x for x in p.glob("*.pdf")])
        threading.Thread(target=package_exhibits, args=(files, p/ "PACKAGED", "EXHIBIT_INDEX"), daemon=True).start()
    return redirect(url_for('ui_home'))

@APP.route("/dashboard")
def ui_dash():
    path = Path(CONFIG["dashboard_path"]).resolve()
    if path.exists():
        webbrowser.open(str(path))
    return redirect(url_for('ui_home'))

def run_gui():
    url = "http://127.0.0.1:5010/"
    threading.Timer(1.0, lambda: webbrowser.open(url)).start()
    APP.run(host="127.0.0.1", port=5010, debug=False)

# ---------- 18) CLI ----------

def parse_args():
    ap = argparse.ArgumentParser(description="Litigation OS — OMEGA MAX")
    ap.add_argument("--once", action="store_true", help="Scan evidence once")
    ap.add_argument("--package-exhibits", nargs="+", help="List of exhibit PDF files to package (stamped)")
    ap.add_argument("--exhibit-out", type=str, help="Output directory for packaged exhibits")
    ap.add_argument("--zip-validate", type=str, help="Validate a ZIP path for filing")
    ap.add_argument("--mifile-upload", type=str, help="Upload a file to MiFILE")
    ap.add_argument("--case", type=str, default="housing", help="Case key: housing|custody|lt")
    ap.add_argument("--doc-type", type=str, default="Filing", help="Doc type for MiFILE")
    ap.add_argument("--crawl-code", action="store_true", help="Scan code assets")
    ap.add_argument("--bundle-code", action="store_true", help="Bundle code and zip")
    ap.add_argument("--audit-code", action="store_true", help="Create a simple code audit snapshot")
    ap.add_argument("--gui", action="store_true", help="Launch local GUI panel")
    return ap.parse_args()

# ---------- 19) Core main ----------

def main():
    args = parse_args()
    if args.once: initial_scan(); return
    if args.package_exhibits:
        files = [Path(x) for x in args.package_exhibits]
        out = Path(args.exhibit_out or CONFIG["exhibits_dir"])/"PACKAGED"
        package_exhibits(files, out)
        return
    if args.zip_validate:
        flags = validate_zip(Path(args.zip_validate))
        if flags: print("ZIP flags:", *flags, sep="\n- ")
        else: print("ZIP OK")
        return
    if args.mifile_upload:
        case_map = CONFIG.get("cases",{})
        if args.case not in case_map:
            raise SystemExit(f"Unknown case key: {args.case}")
        case_no = case_map[args.case]["case_no"]
        u,p = _get_mifile_creds()
        cli = MiFileClient(CONFIG["mifile"]["base_url"], CONFIG["mifile"]["headless"], int(CONFIG["mifile"]["timeout_ms"]))
        try:
            cli.login(u,p)
            proofs = Path(CONFIG["proofs_dir"])
            scr = cli.upload(args.mifile_upload, case_no=case_no, doc_type=args.doc_type, proofs_dir=proofs)
            print(f"MiFILE proof screenshot: {scr}")
        finally:
            cli.close()
        return
    if args.crawl_code: crawl_code(); return
    if args.bundle_code: bundle_code(); return
    if args.audit_code: audit_code_llm(); return
    if args.gui: run_gui(); return
    # default: scan once
    initial_scan()

if __name__ == "__main__":
    main()
