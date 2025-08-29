import os
import json
import zipfile
import hashlib
import re
from typing import Dict, List, TypedDict
from docx import Document

BASE_DIR = os.getenv("LEGAL_RESULTS_DIR", os.path.join("F:/", "LegalResults"))
OUTPUT_DIR = os.path.join(BASE_DIR, "FOIA")
LOG_PATH = os.path.join(OUTPUT_DIR, "foia_request_log.json")

EXPECTED_HASHES: Dict[str, str] = {
    "EGLE_request.docx": "69c71f96efede98c419a4d45465af2cb16518ac6522b7c89904c0f99b52edd2a",
    "County_Clerk_request.docx": "43aae47f5fdc2b6eab8691f1da050ee5fb643aa5f80f496d999c1cbaca84ab8e",
    "Sheriff_request.docx": "d7b5edebbec9624854f6ee57d84b7f8ef758fbce71018e62649f0b3d71c1abd3",
}


class RequestInfo(TypedDict):
    filename: str
    body: List[str]


REQUESTS: Dict[str, RequestInfo] = {
    "EGLE": {
        "filename": "EGLE_request.docx",
        "body": [
            "Subject: Sewer leak reports for 1977 Whitehall Rd (Feb–June 2025).",
            "Please provide all records of inspections, violation notices, and communications with the park owners.",
        ],
    },
    "County_Clerk": {
        "filename": "County_Clerk_request.docx",
        "body": [
            "Subject: Policies on counterclaims in summary proceedings.",
            'Please provide any rejection notices or memos related to filings labelled "moot" since May 2025.',
        ],
    },
    "Sheriff": {
        "filename": "Sheriff_request.docx",
        "body": [
            "Subject: Execution of eviction writ for Lot 17, 1977 Whitehall Rd.",
            "Request any logs or bodycam footage documenting the writ execution.",
        ],
    },
}


def _docx_sha256(path: str) -> str:
    with zipfile.ZipFile(path, "r") as z:
        xml = z.read("word/document.xml")
    xml = re.sub(rb'w:rsid[^\"]*="[^"]*"', b"", xml)
    return hashlib.sha256(xml).hexdigest()


def build_requests() -> None:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    log: List[Dict[str, str]] = []
    hashes: Dict[str, str] = {}
    for key, info in REQUESTS.items():
        doc = Document()
        doc.add_heading(f"FOIA Request – {key}", 0)
        for para in info["body"]:
            doc.add_paragraph(para)
        out_path = os.path.join(OUTPUT_DIR, info["filename"])
        doc.save(out_path)
        file_hash = _docx_sha256(out_path)
        hashes[info["filename"]] = file_hash
        log.append({"agency": key, "file": info["filename"], "hash": file_hash})
    with open(LOG_PATH, "w") as f:
        json.dump(log, f, indent=2)
    zip_path = os.path.join(OUTPUT_DIR, "FOIA_PACKET_SHADY_OAKS_2025.zip")
    with zipfile.ZipFile(zip_path, "w") as z:
        for item in log:
            z.write(os.path.join(OUTPUT_DIR, item["file"]), item["file"])
    with zipfile.ZipFile(zip_path, "r") as z:
        missing = [name for name in EXPECTED_HASHES if name not in z.namelist()]
        if missing:
            raise ValueError(f"Missing files in zip: {', '.join(missing)}")
    for filename, expected in EXPECTED_HASHES.items():
        if hashes.get(filename) != expected:
            raise ValueError(f"Hash mismatch for {filename}")
    print(f"FOIA packet created at {zip_path}")


if __name__ == "__main__":
    build_requests()
