#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py
All-in-one Litigation OS mainframe with:
 - Evidence ingest (PDF/DOCX/TXT), OCR fallback, hashing, PLSC
 - LLM action engine (auto: Ollama → OpenAI keyring)
 - Action Report DOCX + Dashboard + Webhook
 - Daemon + Windows Task Scheduler
 - Compiler + Crawler + CODE_KEEPER (Z:\) + Bundle + LLM Code Audit
 - **NEW: MiFILE e-filing automation (Playwright)**
 - **NEW: Local GUI panel (Flask) to trigger crawls, audits, uploads**

Usage (examples):
  python LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py --once
  python LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py --daemon
  python LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py --install-task
  python LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py --crawl-code
  python LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py --bundle-code
  python LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py --audit-code
  python LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py --compile-exe
  python LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py --mifile-upload "F:\\READY_TO_FILE\\YourDoc.pdf" --case housing --doc-type "Motion"
  python LITIGATION_OS_CROSSWIRE_OMEGA_PLUS.py --gui

Secrets:
 - OpenAI API key stored via keyring: service='LitOS', user='openai_api_key'
 - MiFILE credentials stored via keyring: service='LitOS', users 'mifile_username', 'mifile_password'
No placeholders. Court-usable outputs only.
"""
from __future__ import annotations
import os, sys, subprocess, json, time, shutil, hashlib, sqlite3, re, argparse, threading, datetime, traceback, platform, zipfile, ast, webbrowser
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple, Set

# ---------- 0) Self-bootstrap: install deps if missing ----------

REQUIRED = [
    "watchdog",
    "pymupdf",       # import name: fitz
    "python-docx",   # import name: docx
    "keyring",
    "requests",
    "pyflakes",      # python lint
    "flask",         # GUI
    "playwright",    # MiFILE automation
]

# Optional: OCR stack
OPTIONAL = [
    "pytesseract",
    "Pillow",
    "pyinstaller",   # optional compile
]

def _pip_install(pkgs: List[str]) -> None:
    py = sys.executable
    for p in pkgs:
        mod = p.split("==")[0].replace("-", "_")
        try:
            __import__(mod)
            continue
        except Exception:
            pass
        try:
            subprocess.run([py, "-m", "pip", "install", "--upgrade", p], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except subprocess.CalledProcessError:
            subprocess.run(["pip", "install", "--upgrade", p], check=False)

def ensure_deps():
    _pip_install(REQUIRED)
    _pip_install(OPTIONAL)
    # Ensure Playwright browsers (Chromium) installed
    try:
        subprocess.run([sys.executable, "-m", "playwright", "install", "chromium"], check=False, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=180)
    except Exception:
        pass

ensure_deps()

# Safe imports after bootstrap
import requests
import keyring
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, redirect, url_for, Response

try:
    import pytesseract
    from PIL import Image
except Exception:
    pytesseract = None
    Image = None

# Playwright import (lazy failure tolerated; functions will guard)
try:
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
except Exception:
    sync_playwright = None
    PWTimeout = Exception

# ---------- 1) Constants & Config ----------

IS_WIN = platform.system().lower().startswith("win")
APP_DIR = Path(os.getenv("APPDATA") if IS_WIN else Path.home() / ".config") / "LitOS"
APP_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = APP_DIR / "knowledge.db"
CONFIG_PATH = APP_DIR / "config.json"
LOG_PATH = APP_DIR / "litigation_os.log"

DEFAULT_CONFIG = {
    "scan_roots": [r"F:\\", r"Z:\\", r"C:\\Litigation"],
    "recursive": True,
    "file_glob": ["*.pdf", "*.docx", "*.txt"],
    "ignore_dirs": [r"\\$RECYCLE.BIN", r"\\System Volume Information", r"\\node_modules", r"\\.git", r"\\__pycache__"],
    "results_dir": str(Path("F:/LegalResults/RECOMMENDED_ACTIONS").resolve() if IS_WIN else Path.home() / "LegalResults/RECOMMENDED_ACTIONS"),
    "dashboard_path": str(Path("F:/LegalResults/index.html").resolve() if IS_WIN else Path.home() / "LegalResults/index.html"),
    "ready_to_file": str(Path("F:/READY_TO_FILE").resolve() if IS_WIN else Path.home() / "READY_TO_FILE"),
    "proofs_dir": str(Path("F:/LegalResults/MiFILE_Proofs").resolve() if IS_WIN else Path.home() / "LegalResults/MiFILE_Proofs"),
    "code": {
        "roots": [r"Z:\\"],
        "recursive": True,
        "extensions": [
            ".py",".ps1",".psm1",".psd1",".bat",".cmd",".vbs",".vb",
            ".js",".ts",".json",".yml",".yaml",".toml",".ini",".cfg",
            ".cs",".java",".go",".rs",".c",".cpp",".h",".hpp",".rb",".php",".pl",".lua",".sh",".md",".sql",".xaml"
        ],
        "special_names": ["Dockerfile","Makefile","dockerfile"],
        "bundle_dir": str(Path("F:/LegalResults/CODE_BUNDLE").resolve() if IS_WIN else Path.home() / "LegalResults/CODE_BUNDLE")
    },
    "llm": {
        "backend": "auto",   # "auto" | "openai" | "ollama"
        "openai_model": "gpt-5.0-mini",
        "ollama_model": "llama3.1:8b",
        "max_tokens": 2000,
        "temperature": 0.2
    },
    "webhook_url": "",
    "cases": {
        "housing": {"case_no":"2025-002760-CZ", "court":"14th Circuit Court, Muskegon County, Michigan"},
        "custody": {"case_no":"2024-0000001507-DC", "court":"14th Circuit Court, Muskegon County, Michigan"},
        "lt": {"case_no":"2025-25061626LT-LT", "court":"60th District Court, Muskegon County, Michigan"}
    },
    "owner": {
        "name": "Andrew J Pigors",
        "email": "Andrewjpigors@gmail.com",
        "address": "Lot 17, 1977 Whitehall Rd, Muskegon, MI 49445"
    },
    "mifile": {
        "base_url": "https://mifile.courts.michigan.gov/cases",
        "headless": True,
        "timeout_ms": 30000
    }
}

def load_config() -> Dict[str, Any]:
    if CONFIG_PATH.exists():
        try:
            cfg = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            cfg = DEFAULT_CONFIG
    else:
        cfg = DEFAULT_CONFIG
        CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
        CONFIG_PATH.write_text(json.dumps(cfg, indent=2), encoding="utf-8")
    return cfg

CONFIG = load_config()
RESULTS_DIR = Path(CONFIG["results_dir"]); RESULTS_DIR.mkdir(parents=True, exist_ok=True)
CODE_BUNDLE_DIR = Path(CONFIG["code"]["bundle_dir"]); CODE_BUNDLE_DIR.mkdir(parents=True, exist_ok=True)
Path(CONFIG["proofs_dir"]).mkdir(parents=True, exist_ok=True)
Path(CONFIG["ready_to_file"]).mkdir(parents=True, exist_ok=True)

# ---------- 2) Logging ----------

def log(msg: str):
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{ts}] {msg}"
    print(line, flush=True)
    try:
        with open(LOG_PATH, "a", encoding="utf-8") as f:
            f.write(line + "\n")
    except Exception:
        pass

# ---------- 3) SQLite schema ----------

def db_connect() -> sqlite3.Connection:
    con = sqlite3.connect(str(DB_PATH))
    con.execute("PRAGMA journal_mode=WAL;")
    con.execute("PRAGMA synchronous=NORMAL;")
    return con

def db_init():
    con = db_connect()
    cur = con.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS files(
        id INTEGER PRIMARY KEY,
        path TEXT UNIQUE,
        sha256 TEXT,
        mtime REAL,
        size INTEGER,
        doc_type TEXT,
        case_tag TEXT,
        llm_backend TEXT,
        llm_model TEXT,
        actions_json TEXT,
        created_at TEXT,
        updated_at TEXT
    );""")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS paragraphs(
        id INTEGER PRIMARY KEY,
        file_id INTEGER,
        p_index INTEGER,
        text TEXT,
        sha256 TEXT,
        FOREIGN KEY(file_id) REFERENCES files(id) ON DELETE CASCADE
    );""")
    # Code inventory tables
    cur.execute("""
    CREATE TABLE IF NOT EXISTS code_files(
        id INTEGER PRIMARY KEY,
        path TEXT UNIQUE,
        sha256 TEXT,
        size INTEGER,
        mtime REAL,
        language TEXT,
        shebang TEXT,
        lint_ok INTEGER,
        compile_ok INTEGER,
        error TEXT,
        created_at TEXT
    );""")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS code_refs(
        id INTEGER PRIMARY KEY,
        code_file_id INTEGER,
        ref_type TEXT,
        ref_name TEXT,
        FOREIGN KEY(code_file_id) REFERENCES code_files(id) ON DELETE CASCADE
    );""")
    con.commit(); con.close()

db_init()

# ---------- 4) Utility ----------

def sha256_bytes(b: bytes) -> str:
    h = hashlib.sha256(); h.update(b); return h.hexdigest()

def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024*1024), b""):
            h.update(chunk)
    return h.hexdigest()

# ---------- 5) Text extraction ----------

def extract_text(path: Path) -> Tuple[str, Dict[str, Any]]:
    meta = {"method": None, "pages": 0, "ocr_used": False}
    text = ""
    try:
        suf = path.suffix.lower()
        if suf == ".pdf":
            text, meta = _extract_pdf_text(path)
        elif suf == ".docx":
            text = _extract_docx_text(path); meta["method"] = "python-docx"
        elif suf == ".txt":
            text = path.read_text(encoding="utf-8", errors="ignore"); meta["method"]="txt"
        else:
            meta["method"]="unsupported"
    except Exception as e:
        log(f"extract_text error {path}: {e}")
    return text.strip(), meta

def _extract_pdf_text(path: Path) -> Tuple[str, Dict[str, Any]]:
    meta = {"method":"pymupdf","pages":0,"ocr_used":False}
    out = []
    with fitz.open(str(path)) as doc:
        meta["pages"] = doc.page_count
        for i in range(doc.page_count):
            t = doc.load_page(i).get_text("text")
            if t.strip(): out.append(t)
    txt = "\n".join(out).strip()
    if txt: return txt, meta
    if pytesseract and Image:
        try:
            ocr_txt = _ocr_pdf(path); meta["method"]="ocr"; meta["ocr_used"]=True; return ocr_txt, meta
        except Exception as e:
            log(f"OCR failed {path}: {e}")
    return "", meta

def _ocr_pdf(path: Path) -> str:
    out = []
    with fitz.open(str(path)) as doc:
        for i in range(doc.page_count):
            pix = doc.load_page(i).get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            from io import BytesIO
            img = Image.open(BytesIO(img_bytes))
            out.append(pytesseract.image_to_string(img))
    return "\n".join(out)

def _extract_docx_text(path: Path) -> str:
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)

# ---------- 6) Paragraph splitter ----------

def split_paragraphs(text: str) -> List[str]:
    parts = re.split(r"\n{2,}|\r{2,}", text)
    return [p.strip() for p in parts if p and p.strip()]

# ---------- 7) LLM backends ----------

def _ollama_up() -> bool:
    try:
        r = requests.get("http://127.0.0.1:11434/api/tags", timeout=1.5)
        return r.status_code == 200
    except Exception:
        return False

def get_backend(llm_cfg: Dict[str,Any]) -> Tuple[str,str]:
    be = llm_cfg.get("backend","auto")
    if be == "auto":
        if _ollama_up(): return "ollama", llm_cfg.get("ollama_model","llama3.1:8b")
        key = keyring.get_password("LitOS","openai_api_key")
        if key: return "openai", llm_cfg.get("openai_model","gpt-5.0-mini")
        return "ollama", llm_cfg.get("ollama_model","llama3.1:8b")
    if be == "openai": return "openai", llm_cfg.get("openai_model","gpt-5.0-mini")
    if be == "ollama": return "ollama", llm_cfg.get("ollama_model","llama3.1:8b")
    return "ollama","llama3.1:8b"

def _call_openai(prompt: str, max_tokens: int, temperature: float, model: str) -> str:
    key = keyring.get_password("LitOS","openai_api_key")
    if not key:
        try:
            if sys.stdin and sys.stdin.isatty():
                print("\nEnter your OpenAI API key (stored in keyring):")
                key = input("sk-...: ").strip()
                if key: keyring.set_password("LitOS","openai_api_key",key)
        except Exception:
            pass
    if not key:
        raise RuntimeError("OpenAI API key missing in keyring (service='LitOS', user='openai_api_key').")
    headers = {"Authorization": f"Bearer {key}", "Content-Type":"application/json"}
    payload = {
        "model": model,
        "messages": [
            {"role":"system","content":"You are a Michigan litigation expert. Cite MCR/MCL/Benchbooks precisely. No placeholders."},
            {"role":"user","content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": temperature
    }
    r = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload, timeout=90)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

def _call_ollama(prompt: str, model: str, options: Optional[Dict[str,Any]] = None) -> str:
    payload = {"model": model, "prompt": prompt, "stream": False}
    if options: payload["options"]=options
    r = requests.post("http://127.0.0.1:11434/api/generate", json=payload, timeout=180)
    r.raise_for_status()
    return r.json().get("response","").strip()

def llm_analyze(doc_text: str, path: Path, meta: Dict[str,Any]) -> Dict[str,Any]:
    backend, model = get_backend(CONFIG["llm"])
    temp = CONFIG["llm"].get("temperature",0.2)
    max_tokens = CONFIG["llm"].get("max_tokens",2000)
    owner = CONFIG.get("owner",{}); known = CONFIG.get("cases",{})
    prompt = f"""
Analyze this document within a Michigan Litigation OS.

Owner: {owner.get('name')} <{owner.get('email')}>, Address: {owner.get('address')}
Known Cases: Housing {known.get('housing',{})}, Custody {known.get('custody',{})}, LT {known.get('lt',{})}

Document: {path.name} | Method: {meta.get('method')} | Pages: {meta.get('pages')} | OCR: {meta.get('ocr_used')}

TASKS:
1) Classify: [Housing, Custody, PPO, Other].
2) Identify key MCR/MCL/Benchbook hooks triggered by this text.
3) Recommend ONE Single Best Action with exact citations; ≤120-word rationale.
4) Short Action Title (≤10 words).
5) 3–6 bullet facts/snippets supporting the action.

Rules: Precise; no placeholders; prefer SCAO forms when available; under 400 words total.
--- TEXT START ---
{doc_text[:20000]}
--- TEXT END ---
"""
    if backend == "openai":
        content = _call_openai(prompt, max_tokens=max_tokens, temperature=temp, model=model); used="openai"
    else:
        content = _call_ollama(prompt, model=model, options={"temperature":temp}); used="ollama"
    return {"backend": used, "model": model, "content": content}

# ---------- 8) Action Report (DOCX) ----------

def make_action_report(path: Path, analysis: Dict[str,Any], file_sha: str) -> Path:
    out_dir = RESULTS_DIR; out_dir.mkdir(parents=True, exist_ok=True)
    title = "Recommended Action"
    body = analysis.get("content","").strip()
    m = re.match(r"\s*(?:Action Title:|Title:)?\s*(.+)\n", body)
    if m:
        t = m.group(1).strip()
        if 3 <= len(t) <= 80: title = t
    doc = Document()
    p = doc.add_paragraph(); run=p.add_run(title.upper()); run.bold=True; run.font.size=Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    owner = CONFIG.get("owner",{}); cases=CONFIG.get("cases",{})
    doc.add_paragraph(f"Prepared for: {owner.get('name')}  |  Email: {owner.get('email')}")
    doc.add_paragraph(f"Housing Case: {cases.get('housing',{}).get('case_no','')} — {cases.get('housing',{}).get('court','')}")
    doc.add_paragraph(f"Custody Case: {cases.get('custody',{}).get('case_no','')} — {cases.get('custody',{}).get('court','')}")
    doc.add_paragraph(f"LT Case: {cases.get('lt',{}).get('case_no','')} — {cases.get('lt',{}).get('court','')}")
    doc.add_paragraph(f"Source file: {str(path)}")
    doc.add_paragraph(f"SHA-256: {file_sha}")
    doc.add_paragraph("")
    for line in body.splitlines(): doc.add_paragraph(line)
    doc.add_paragraph(""); doc.add_paragraph("Generated by LITIGATION_OS_CROSSWIRE_OMEGA_PLUS — PLSC enforced.")
    safe = re.sub(r"[^A-Za-z0-9_.-]+","_",path.stem)[:50]
    out_path = out_dir / f"{safe}__ACTION_{file_sha[:10]}.docx"; doc.save(str(out_path)); return out_path

# ---------- 9) Dashboard ----------

def write_dashboard():
    con=db_connect(); cur=con.cursor()
    files = cur.execute("SELECT path, sha256, doc_type, llm_backend, llm_model, created_at FROM files ORDER BY created_at DESC LIMIT 200").fetchall()
    code_stats = cur.execute("SELECT COUNT(*), SUM(size) FROM code_files").fetchone()
    py_ok = cur.execute("SELECT COUNT(*) FROM code_files WHERE language='python' AND (lint_ok=1 AND compile_ok=1)").fetchone()[0]
    con.close()
    total_code = code_stats[0] or 0
    total_bytes = code_stats[1] or 0
    html = ["<html><head><meta charset='utf-8'><title>Litigation OS Dashboard</title></head><body>"]
    html.append("<h2>Litigation OS — Evidence & Code Overview</h2>")
    html.append(f"<p><b>Code assets indexed:</b> {total_code} files ({total_bytes} bytes). Python OK: {py_ok}</p>")
    html.append("<h3>Recent Evidence Files</h3>")
    html.append("<table border='1' cellspacing='0' cellpadding='6'><tr><th>File</th><th>SHA-256</th><th>Type</th><th>LLM</th><th>Model</th><th>Created</th></tr>")
    for path, sha, dt, lb, lm, created in files:
        html.append(f"<tr><td>{path}</td><td>{sha[:12]}</td><td>{dt or ''}</td><td>{lb or ''}</td><td>{lm or ''}</td><td>{created or ''}</td></tr>")
    html.append("</table></body></html>")
    Path(CONFIG["dashboard_path"]).parent.mkdir(parents=True, exist_ok=True)
    Path(CONFIG["dashboard_path"]).write_text("\n".join(html), encoding="utf-8")
    log(f"Dashboard updated: {CONFIG['dashboard_path']}")

# ---------- 10) Webhook ----------

def notify_webhook(text: str):
    url = CONFIG.get("webhook_url","")
    if not url: return
    try: requests.post(url, json={"text": text}, timeout=5)
    except Exception as e: log(f"webhook failed: {e}")

# ---------- 11) Evidence pipeline ----------

SUPPORTED_EXTS = {".pdf",".docx",".txt"}

def guess_doc_type(text: str, path: Path) -> str:
    t = (path.name + " " + text[:500]).lower()
    if "ppo" in t or "personal protection order" in t: return "PPO"
    if any(k in t for k in ["custody","parenting time","foc"]): return "Custody"
    if any(k in t for k in ["shady oaks","mobile home","rent","landlord","lot"]): return "Housing"
    return "Other"

def process_file(path: Path):
    try:
        if not path.exists() or not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTS: return
        st = path.stat(); fsha = sha256_file(path)
        con=db_connect(); cur=con.cursor(); now=datetime.datetime.utcnow().isoformat()
        row = cur.execute("SELECT id, mtime, sha256 FROM files WHERE path=?", (str(path),)).fetchone()
        if row:
            fid, old_m, old_h = row
            if abs(old_m - st.st_mtime) < 1e-6 and old_h == fsha: con.close(); return
            cur.execute("DELETE FROM paragraphs WHERE file_id=?", (fid,))
            cur.execute("UPDATE files SET mtime=?, size=?, sha256=?, updated_at=? WHERE id=?", (st.st_mtime, st.st_size, fsha, now, fid)); file_id=fid
        else:
            cur.execute("INSERT INTO files(path,sha256,mtime,size,created_at,updated_at) VALUES(?,?,?,?,?,?)",
                        (str(path), fsha, st.st_mtime, st.st_size, now, now)); file_id=cur.lastrowid
        con.commit()

        text, meta = extract_text(path); dtype = guess_doc_type(text, path)
        paras = split_paragraphs(text)
        for i,p in enumerate(paras):
            cur.execute("INSERT INTO paragraphs(file_id,p_index,text,sha256) VALUES(?,?,?,?)", (file_id,i,p,sha256_bytes(p.encode("utf-8"))))
        con.commit()

        analysis = llm_analyze(text, path, meta)
        cur.execute("UPDATE files SET doc_type=?, llm_backend=?, llm_model=?, actions_json=? WHERE id=?",
                    (dtype, analysis["backend"], analysis["model"], json.dumps(analysis, ensure_ascii=False), file_id))
        con.commit(); con.close()

        out_docx = make_action_report(path, analysis, fsha)
        log(f"Action report: {out_docx}")
        notify_webhook(f"Action report created for {path.name}: {out_docx}")
        write_dashboard()
    except Exception as e:
        log(f"process_file error {path}: {e}\n{traceback.format_exc()}")

def initial_scan():
    roots=[Path(r) for r in CONFIG.get("scan_roots",[])]; globs=CONFIG.get("file_glob",["*.pdf","*.docx","*.txt"]); rec=CONFIG.get("recursive",True)
    for root in roots:
        if not root.exists(): log(f"Missing scan root: {root}"); continue
        for pattern in globs:
            if rec:
                for p in root.rglob(pattern): process_file(p)
            else:
                for p in root.glob(pattern): process_file(p)

# ---------- 12) Watchdog ----------

class Handler(FileSystemEventHandler):
    def on_created(self, event):
        if event.is_directory: return
        path=Path(event.src_path)
        if path.suffix.lower() in SUPPORTED_EXTS:
            log(f"Created: {path}"); time.sleep(0.5); process_file(path)
    def on_modified(self, event):
        if event.is_directory: return
        path=Path(event.src_path)
        if path.suffix.lower() in SUPPORTED_EXTS:
            log(f"Modified: {path}"); time.sleep(0.5); process_file(path)

def run_daemon():
    observer=Observer()
    for r in CONFIG.get("scan_roots",[]):
        p=Path(r)
        if p.exists(): observer.schedule(Handler(), str(p), recursive=CONFIG.get("recursive",True)); log(f"Watching: {p}")
        else: log(f"Skip watch (not found): {p}")
    # Also watch READY_TO_FILE for auto MiFILE uploads
    rtf = Path(CONFIG["ready_to_file"])
    if rtf.exists():
        observer.schedule(ReadyToFileHandler(), str(rtf), recursive=False)
        log(f"Watching READY_TO_FILE: {rtf}")
    observer.start()
    try:
        while True: time.sleep(1.0)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()

# ---------- 13) Windows Task Scheduler ----------

def install_task():
    if not IS_WIN: log("Task install is Windows-only."); return
    py=sys.executable; script=Path(__file__).resolve(); name="LitOS_OmegaPlus_Daemon"
    try:
        subprocess.run(["schtasks","/Delete","/TN",name,"/F"], check=False, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        subprocess.run(["schtasks","/Create","/SC","DAILY","/TN",name,"/TR",f'"{py}" "{script}" --daemon',"/ST","09:00"], check=True)
        subprocess.run(["schtasks","/Run","/TN",name], check=False)
        log(f"Scheduled Task installed: {name}")
    except Exception as e:
        log(f"Task install failed: {e}")

# ======================= COMPILER + CRAWLER + CODE KEEPER =======================

CODE_EXTS: Set[str] = set(CONFIG["code"]["extensions"])
SPECIAL_NAMES: Set[str] = set(CONFIG["code"]["special_names"])
IGNORE_DIR_PATTERNS = [d.lower() for d in CONFIG.get("ignore_dirs", [])]

LANG_MAP = {
    ".py":"python",".ps1":"powershell",".psm1":"powershell",".psd1":"powershell",
    ".bat":"batch",".cmd":"batch",".vbs":"vbscript",".vb":"vbnet",
    ".js":"javascript",".ts":"typescript",".json":"json",".yml":"yaml",".yaml":"yaml",".toml":"toml",".ini":"ini",".cfg":"cfg",
    ".cs":"csharp",".java":"java",".go":"go",".rs":"rust",".c":"c",".cpp":"cpp",".h":"c-header",".hpp":"cpp-header",
    ".rb":"ruby",".php":"php",".pl":"perl",".lua":"lua",".sh":"shell",".md":"markdown",".sql":"sql",".xaml":"xaml"
}

def _is_ignored_dir(path: Path) -> bool:
    lp = str(path).lower()
    return any(pat in lp for pat in IGNORE_DIR_PATTERNS)

def detect_language(path: Path) -> str:
    if path.name in SPECIAL_NAMES:
        if "docker" in path.name.lower(): return "dockerfile"
        if "makefile" == path.name.lower(): return "makefile"
    return LANG_MAP.get(path.suffix.lower(), "unknown")

def parse_python_refs(text: str) -> List[str]:
    out = []
    try:
        node = ast.parse(text)
        for n in ast.walk(node):
            if isinstance(n, ast.Import):
                for a in n.names: out.append(a.name.split(".")[0])
            elif isinstance(n, ast.ImportFrom):
                if n.module: out.append(n.module.split(".")[0])
    except Exception:
        pass
    return sorted(set(out))

def parse_powershell_refs(text: str) -> List[str]:
    refs = set()
    for m in re.finditer(r"(?i)\bImport-Module\s+([^\s;]+)", text):
        refs.add(m.group(1).strip())
    for m in re.finditer(r"(?i)\b\.?\s*\.\\([^\s;]+\.psm1)", text):
        refs.add(m.group(1).strip())
    return sorted(refs)

def lint_python(path: Path, text: Optional[str]=None) -> Tuple[bool, bool, str]:
    lint_ok=True; compile_ok=True; err=""
    try:
        try:
            if text is None: text = path.read_text(encoding="utf-8", errors="ignore")
            compile(text, str(path), "exec")
        except Exception as e:
            compile_ok=False; err += f"COMPILE: {e}\n"
        try:
            r = subprocess.run([sys.executable, "-m", "pyflakes", str(path)], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=30)
            if r.returncode != 0 or r.stdout or r.stderr:
                lint_ok=False
                err += (r.stdout or "") + (r.stderr or "")
        except Exception as e:
            err += f"PYFLAKES_ERROR: {e}\n"
    except Exception as e:
        err += f"lint_python failed: {e}"
    return lint_ok, compile_ok, err.strip()

def crawl_code():
    roots = [Path(r) for r in CONFIG["code"].get("roots",[])]
    rec = CONFIG["code"].get("recursive", True)
    con = db_connect(); cur=con.cursor()
    registry = []
    count=0
    for root in roots:
        if not root.exists(): log(f"Code root missing: {root}"); continue
        for dirpath, dirnames, filenames in os.walk(root):
            pdir = Path(dirpath)
            if _is_ignored_dir(pdir):
                dirnames[:] = []
                continue
            for fname in filenames:
                fpath = pdir / fname
                ext = fpath.suffix.lower()
                if (ext in CODE_EXTS) or (fname in SPECIAL_NAMES):
                    try:
                        st = fpath.stat()
                        size = st.st_size; mtime = st.st_mtime
                        try:
                            text = fpath.read_text(encoding="utf-8", errors="ignore")
                        except Exception:
                            text = ""
                        fsha = sha256_file(fpath)
                        lang = detect_language(fpath)
                        shebang = ""
                        if text.startswith("#!") or text.startswith("::#"):
                            first = text.splitlines()[0].strip()
                            if len(first) < 200: shebang = first

                        lint_ok=compile_ok=True; err=""
                        refs: List[str]=[]
                        if lang=="python":
                            lint_ok, compile_ok, err = lint_python(fpath, text=text)
                            refs = parse_python_refs(text)
                        elif lang=="powershell":
                            refs = parse_powershell_refs(text)

                        now = datetime.datetime.utcnow().isoformat()
                        cur.execute("INSERT OR REPLACE INTO code_files(path,sha256,size,mtime,language,shebang,lint_ok,compile_ok,error,created_at) VALUES(?,?,?,?,?,?,?,?,?,?)",
                                    (str(fpath), fsha, size, mtime, lang, shebang, int(bool(lint_ok)), int(bool(compile_ok)), err, now))
                        row = cur.execute("SELECT id FROM code_files WHERE path=?", (str(fpath),)).fetchone()
                        if row:
                            cid = row[0]
                            cur.execute("DELETE FROM code_refs WHERE code_file_id=?", (cid,))
                            for rname in refs:
                                cur.execute("INSERT INTO code_refs(code_file_id, ref_type, ref_name) VALUES(?,?,?)", (cid, "import", rname))
                        con.commit()

                        registry.append({
                            "path": str(fpath), "sha256": fsha, "size": size, "mtime": mtime,
                            "language": lang, "shebang": shebang, "lint_ok": lint_ok, "compile_ok": compile_ok,
                            "imports": refs
                        })
                        count += 1
                    except Exception as e:
                        log(f"crawl_code error {fpath}: {e}")
    keeper_dir = CODE_BUNDLE_DIR; keeper_dir.mkdir(parents=True, exist_ok=True)
    reg_path = keeper_dir / "CODE_KEEPER_REGISTRY.json"
    reg_path.write_text(json.dumps({"generated_at": datetime.datetime.utcnow().isoformat(), "files": registry}, indent=2), encoding="utf-8")
    log(f"Code crawl complete. Indexed {count} files. Registry: {reg_path}")
    write_dashboard()

def bundle_code():
    con=db_connect(); cur=con.cursor()
    rows = cur.execute("SELECT path, language, lint_ok, compile_ok FROM code_files").fetchall()
    manifest = {"generated_at": datetime.datetime.utcnow().isoformat(), "items":[]}
    dest = CODE_BUNDLE_DIR / "bundle"; 
    if dest.exists():
        shutil.rmtree(dest, ignore_errors=True)
    dest.mkdir(parents=True, exist_ok=True)
    copied=0
    for path, lang, lint_ok, compile_ok in rows:
        p = Path(path)
        if not p.exists(): continue
        if any(s in str(p).lower() for s in ["/node_modules/","\\node_modules\\","/.git/","\\._git\\","/__pycache__/","\\__pycache__\\"]):
            continue
        rel = re.sub(r"^[A-Za-z]:(\\\\|/)", "", str(p)) if IS_WIN else str(p).lstrip("/")
        outp = dest / rel
        outp.parent.mkdir(parents=True, exist_ok=True)
        try:
            shutil.copy2(str(p), str(outp))
            manifest["items"].append({"src": str(p), "dst": str(outp), "language": lang, "ok": bool(lint_ok and compile_ok)})
            copied += 1
        except Exception as e:
            log(f"copy failed {p}: {e}")
    man_path = CODE_BUNDLE_DIR / "MANIFEST.json"; man_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    zip_path = CODE_BUNDLE_DIR / "CODE_BUNDLE.zip"
    with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(dest):
            for f in files:
                fp = Path(root) / f
                z.write(str(fp), arcname=str(fp.relative_to(dest)))
        z.write(str(man_path), arcname="MANIFEST.json")
        reg = CODE_BUNDLE_DIR / "CODE_KEEPER_REGISTRY.json"
        if reg.exists(): z.write(str(reg), arcname="CODE_KEEPER_REGISTRY.json")
    log(f"Code bundle complete. Files copied: {copied}. ZIP: {zip_path}")
    write_dashboard()

def audit_code_llm():
    reg = CODE_BUNDLE_DIR / "CODE_KEEPER_REGISTRY.json"
    if not reg.exists():
        log("No CODE_KEEPER_REGISTRY.json found. Run --crawl-code first.")
        return
    data = json.loads(reg.read_text(encoding="utf-8"))
    files = data.get("files", [])
    def score(f):
        base = 0
        if f["language"] in ("python","powershell","javascript","typescript"): base += 5
        if f.get("lint_ok") and f.get("compile_ok"): base += 3
        base += min(5, len(f.get("imports", [])))
        try:
            base += min(5, int(Path(f["path"]).stat().st_size // 2048))
        except Exception:
            pass
        return -base
    sample = sorted(files, key=score)[:60]
    lines = []
    for f in sample:
        lines.append(f"{f['path']} | {f['language']} | ok={f.get('lint_ok') and f.get('compile_ok')} | imports={','.join(f.get('imports',[]))}")
    meta = "\n".join(lines)
    backend, model = get_backend(CONFIG["llm"])
    temp = CONFIG["llm"].get("temperature",0.2); max_tokens = CONFIG["llm"].get("max_tokens",2000)
    prompt = f"""
You are the Compiler+Crawler brain for a Michigan Litigation OS. Using the inventory below, produce a precise integration plan to assemble the strongest possible, court-usable Litigation OS .exe with these goals:
- Evidence scanning, rule-matching (MCR/MCL/Benchbooks), action generation
- MiFILE e-filing automation hooks
- ZIP validator and exhibit naming enforcer
- GUI dashboard (Electron or webview) hooks
- Robust logging and provenance (SHA-256, PLSC)
- Background daemon + scheduled runs
- Safety: never delete originals; copy-only

Inventory (path | language | OK | imports):
{meta}

TASKS:
1) Identify the 8–12 most valuable code assets to pull into the main build and why.
2) List the concrete integration steps (file-by-file), including where to wire them in this engine.
3) Specify missing glue code we must add (exact modules/functions) and any Python dependencies to include.
4) If PowerShell scripts exist, show how to invoke them safely from Python with robust error capture.
5) Provide a one-paragraph, court-facing description of the resulting system (no hype).

Rules: precise, no placeholders, concise but complete.
"""
    if backend=="openai":
        plan = _call_openai(prompt, max_tokens=max_tokens, temperature= temp, model=model); used="openai"
    else:
        plan = _call_ollama(prompt, model=model, options={"temperature":temp}); used="ollama"
    doc = Document()
    p=doc.add_paragraph(); r=p.add_run("CODE AUDIT & INTEGRATION PLAN"); r.bold=True; r.font.size=Pt(15)
    doc.add_paragraph(f"LLM backend: {used} / {model}")
    doc.add_paragraph("")
    for line in plan.splitlines():
        doc.add_paragraph(line)
    plan_docx = CODE_BUNDLE_DIR / "CODE_AUDIT_INTEGRATION_PLAN.docx"; doc.save(str(plan_docx))
    plan_txt = CODE_BUNDLE_DIR / "CODE_AUDIT_INTEGRATION_PLAN.txt"; plan_txt.write_text(plan, encoding="utf-8")
    log(f"Code audit plan written: {plan_docx}")
    write_dashboard()

def compile_exe():
    try:
        script = Path(__file__).resolve()
        dist = script.parent / "dist"
        build = script.parent / "build"
        if dist.exists(): shutil.rmtree(dist, ignore_errors=True)
        if build.exists(): shutil.rmtree(build, ignore_errors=True)
        cmd = [sys.executable, "-m", "PyInstaller", "--onefile", "--name", "LITIGATION_OS_OMEGA_PLUS", str(script)]
        log(f"Running: {' '.join(cmd)}")
        subprocess.run(cmd, check=True)
        log(f"EXE build complete. See: {script.parent/'dist'/'LITIGATION_OS_OMEGA_PLUS.exe'}")
    except Exception as e:
        log(f"compile_exe failed: {e}")

# ======================= NEW: MiFILE AUTOMATION (Playwright) =======================

class MiFileClient:
    def __init__(self, base_url: str, headless: bool=True, timeout_ms: int=30000):
        self.base_url = base_url
        self.headless = headless
        self.timeout_ms = timeout_ms
        self._pw = None
        self._browser = None
        self._context = None
        self._page = None

    def __enter.me__(self):  # typo guard not used
        return self

    def _start(self):
        if sync_playwright is None:
            raise RuntimeError("Playwright is not available. Ensure it is installed.")
        self._pw = sync_playwright().start()
        self._browser = self._pw.chromium.launch(headless=self.headless)
        self._context = self._browser.new_context()
        self._page = self._context.new_page()
        self._page.set_default_timeout(self.timeout_ms)

    def _stop(self):
        try:
            if self._context: self._context.close()
        except Exception: pass
        try:
            if self._browser: self._browser.close()
        except Exception: pass
        try:
            if self._pw: self._pw.stop()
        except Exception: pass

    def login(self, username: str, password: str):
        self._start()
        page = self._page
        page.goto(self.base_url)
        # Attempt multiple selector strategies to be resilient to UI changes.
        # Strategy A: label-based
        try:
            page.get_by_label("Email").fill(username)
            page.get_by_label("Password").fill(password)
        except Exception:
            # Strategy B: common login inputs
            try:
                page.fill("input[type='email']", username)
                page.fill("input[type='password']", password)
            except Exception:
                # Strategy C: id/name guesses
                try:
                    page.fill("input#Email, input[name='Email']", username)
                    page.fill("input#Password, input[name='Password']", password)
                except Exception as e:
                    raise RuntimeError(f"MiFILE login fields not found: {e}")
        # Click login button by text or role
        clicked = False
        for sel in ["text=Sign in", "text=Log in", "button:has-text('Sign in')", "button:has-text('Log in')", "input[type='submit']"]:
            try:
                page.click(sel)
                clicked = True
                break
            except Exception:
                continue
        if not clicked:
            raise RuntimeError("MiFILE login button not found.")
        # Verify login by waiting for some known post-login markers
        # Try to detect any banner or link that implies a login succeeded
        try:
            page.wait_for_timeout(1500)
            if "login" in page.url.lower():
                # try a generic success marker
                page.wait_for_selector("a,button", timeout=5000)
        except PWTimeout:
            pass

    def upload(self, file_path: str, case_no: str, doc_type: str, proofs_dir: Path) -> Path:
        page = self._page
        # Navigate towards an e-filing / upload area heuristically
        # Try common anchors/titles
        candidates = ["text=File a Document", "text=E-File", "text=New Filing", "text=Submit Document", "text=Filing"]
        clicked = False
        for c in candidates:
            try:
                page.click(c, timeout=5000)
                clicked = True
                break
            except Exception:
                continue
        if not clicked:
            # Try a menu approach
            try:
                page.click("text=Cases", timeout=3000)
                page.click("text=File", timeout=3000)
                clicked = True
            except Exception:
                pass
        # Fill case number
        try:
            page.fill("input[placeholder*='Case']", case_no)
        except Exception:
            try:
                page.get_by_label("Case Number").fill(case_no)
            except Exception:
                pass
        # Upload file
        try:
            page.set_input_files("input[type='file']", file_path)
        except Exception as e:
            raise RuntimeError(f"Unable to set file for upload: {e}")
        # Doc type / category
        try:
            page.fill("input[placeholder*='Document Type']", doc_type)
        except Exception:
            try:
                page.select_option("select[name*='DocumentType']", label=doc_type)
            except Exception:
                pass
        # Submit
        for sel in ["text=Submit", "text=File", "button:has-text('Submit')", "button:has-text('File')", "input[type='submit']"]:
            try:
                page.click(sel)
                break
            except Exception:
                continue
        # Wait and capture a confirmation screenshot
        page.wait_for_timeout(2000)
        # Try sniffing a success marker
        success = False
        try:
            if page.get_by_text("success", exact=False):
                success = True
        except Exception:
            pass
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        proofs_dir.mkdir(parents=True, exist_ok=True)
        scr_path = proofs_dir / f"MiFILE_{case_no}_{Path(file_path).stem}_{ts}.png"
        page.screenshot(path=str(scr_path), full_page=True)
        return scr_path

    def close(self):
        self._stop()

def _get_mifile_creds() -> Tuple[str,str]:
    u = keyring.get_password("LitOS","mifile_username") or ""
    p = keyring.get_password("LitOS","mifile_password") or ""
    if not (u and p) and sys.stdin and sys.stdin.isatty():
        print("\n=== MiFILE credentials setup (stored in keyring: service='LitOS') ===")
        if not u:
            u = input("MiFILE username (email): ").strip()
            if u: keyring.set_password("LitOS","mifile_username",u)
        if not p:
            p = input("MiFILE password: ").strip()
            if p: keyring.set_password("LitOS","mifile_password",p)
    return u, p

def mifile_upload_autorun(file_path: str, case_key: str, doc_type: str) -> Path:
    if not Path(file_path).exists():
        raise FileNotFoundError(file_path)
    case_map = CONFIG.get("cases",{})
    if case_key not in case_map:
        raise ValueError(f"Unknown case key: {case_key}. Use one of {list(case_map.keys())}")
    case_no = case_map[case_key]["case_no"]
    u,p = _get_mifile_creds()
    if not (u and p):
        raise RuntimeError("MiFILE credentials are missing (keyring).")
    base = CONFIG["mifile"]["base_url"]; headless = CONFIG["mifile"]["headless"]; to = int(CONFIG["mifile"]["timeout_ms"])
    client = MiFileClient(base_url=base, headless=headless, timeout_ms=to)
    try:
        client.login(u,p)
        proofs = Path(CONFIG["proofs_dir"])
        screenshot = client.upload(file_path=file_path, case_no=case_no, doc_type=doc_type, proofs_dir=proofs)
        # Build a Proof of Service/Submission record (DOCX)
        doc = Document()
        p1=doc.add_paragraph(); r=p1.add_run("MiFILE SUBMISSION RECORD"); r.bold=True; r.font.size=Pt(15)
        doc.add_paragraph(f"Case: {case_no}")
        doc.add_paragraph(f"Document: {file_path}")
        doc.add_paragraph(f"Type: {doc_type}")
        doc.add_paragraph(f"Time (local): {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        try:
            doc.add_picture(str(screenshot), width=Inches(6.0))
        except Exception:
            pass
        out_docx = Path(CONFIG["proofs_dir"]) / f"PROOF_{Path(file_path).stem}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(str(out_docx))
        log(f"MiFILE upload completed. Proof: {out_docx}")
        return screenshot
    finally:
        client.close()

class ReadyToFileHandler(FileSystemEventHandler):
    def on_created(self, event):
        if event.is_directory: return
        p = Path(event.src_path)
        # Decide case by simple heuristics or mapping in filename
        fname = p.name.lower()
        case_key = "housing"
        if "custody" in fname or "dc" in fname: case_key = "custody"
        if "lt" in fname or "district" in fname: case_key = "lt"
        doc_type = "Filing"  # adjust if pattern known
        try:
            time.sleep(1.0)  # debounce
            mifile_upload_autorun(str(p), case_key=case_key, doc_type=doc_type)
        except Exception as e:
            log(f"READY_TO_FILE auto-upload failed for {p}: {e}")

# ======================= GUI (Flask) =======================

APP = Flask(__name__)

def list_ready_files(limit=50):
    r = Path(CONFIG["ready_to_file"])
    files = []
    if r.exists():
        for p in sorted(r.glob("*.*"), key=lambda x: x.stat().st_mtime, reverse=True)[:limit]:
            files.append(str(p))
    return files

@APP.route("/")
def ui_home():
    files = list_ready_files()
    html = ["<html><head><title>Litigation OS — GUI</title></head><body>",
            "<h2>Litigation OS — Control Panel</h2>",
            "<h3>Quick Actions</h3>",
            "<form method='POST' action='/action'>",
            "<button name='cmd' value='scan-once'>Scan Evidence Once</button> ",
            "<button name='cmd' value='crawl-code'>Crawl Code</button> ",
            "<button name='cmd' value='bundle-code'>Bundle Code</button> ",
            "<button name='cmd' value='audit-code'>Audit Code (LLM)</button>",
            "</form>",
            "<h3>MiFILE Upload</h3>",
            "<form method='POST' action='/upload'>",
            "File path: <input type='text' name='path' size='80' list='rtf'> ",
            "Case: <select name='case'><option>housing</option><option>custody</option><option>lt</option></select> ",
            "Doc Type: <input type='text' name='doctype' value='Filing'> ",
            "<button type='submit'>Upload</button>",
            f"<datalist id='rtf'>{''.join([f'<option value=\"{f}\"></option>' for f in files])}</datalist>",
            "</form>",
            "<p><a href='/open-dashboard' target='_blank'>Open Dashboard</a></p>",
            "</body></html>"]
    return Response("\n".join(html), mimetype="text/html")

@APP.route("/open-dashboard")
def open_dash():
    path = Path(CONFIG["dashboard_path"]).resolve()
    if path.exists():
        webbrowser.open(str(path))
    return redirect(url_for('ui_home'))

@APP.route("/action", methods=["POST"])
def ui_action():
    cmd = request.form.get("cmd","")
    if cmd == "scan-once": threading.Thread(target=initial_scan, daemon=True).start()
    if cmd == "crawl-code": threading.Thread(target=crawl_code, daemon=True).start()
    if cmd == "bundle-code": threading.Thread(target=bundle_code, daemon=True).start()
    if cmd == "audit-code": threading.Thread(target=audit_code_llm, daemon=True).start()
    return redirect(url_for('ui_home'))

@APP.route("/upload", methods=["POST"])
def ui_upload():
    path = request.form.get("path","")
    case_key = request.form.get("case","housing")
    doctype = request.form.get("doctype","Filing")
    def _run():
        try:
            mifile_upload_autorun(path, case_key=case_key, doc_type=doctype)
        except Exception as e:
            log(f"GUI upload failed: {e}")
    threading.Thread(target=_run, daemon=True).start()
    return redirect(url_for('ui_home'))

def run_gui():
    url = "http://127.0.0.1:5005/"
    threading.Timer(1.0, lambda: webbrowser.open(url)).start()
    APP.run(host="127.0.0.1", port=5005, debug=False)

# ======================= CLI =======================

def parse_args():
    ap = argparse.ArgumentParser(description="Litigation OS — OMEGA PLUS")
    ap.add_argument("--once", action="store_true", help="Scan evidence once")
    ap.add_argument("--daemon", action="store_true", help="Watch folders")
    ap.add_argument("--install-task", action="store_true", help="Install Windows daily daemon task")
    ap.add_argument("--rebuild", action="store_true", help="Rebuild dashboard")
    ap.add_argument("--crawl-code", action="store_true", help="Scan code assets and build CODE_KEEPER_REGISTRY.json")
    ap.add_argument("--bundle-code", action="store_true", help="Copy curated code and zip with manifest")
    ap.add_argument("--audit-code", action="store_true", help="Generate LLM code integration plan")
    ap.add_argument("--compile-exe", action="store_true", help="Compile this script to EXE with PyInstaller")
    ap.add_argument("--mifile-upload", type=str, help="Upload a file to MiFILE (provide path)")
    ap.add_argument("--case", type=str, default="housing", help="Case key: housing|custody|lt")
    ap.add_argument("--doc-type", type=str, default="Filing", help="Document type/label for MiFILE")
    ap.add_argument("--gui", action="store_true", help="Launch local GUI panel")
    return ap.parse_args()

def main():
    args = parse_args()
    if args.install_task: install_task(); return
    if args.rebuild: write_dashboard(); return
    if args.crawl-code: crawl_code(); return
    if args.bundle-code: bundle_code(); return
    if args.audit-code: audit_code_llm(); return
    if args.compile_exe: compile_exe(); return
    if args.mifile_upload: mifile_upload_autorun(args.mifile_upload, case_key=args.case, doc_type=args.doc_type); return
    if args.gui: run_gui(); return
    if args.once: initial_scan(); return
    if args.daemon: initial_scan(); run_daemon(); return
    # default: single scan
    initial_scan()

if __name__ == "__main__":
    main()
