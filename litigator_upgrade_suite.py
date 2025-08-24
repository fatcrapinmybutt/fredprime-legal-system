# -*- coding: utf-8 -*-
"""
LITIGATOR UPGRADE SUITE v1 (Windows) — plug-in for GOLDEN LITIGATOR OS
Builds: Exhibit Binder (index + merged PDF), MiFile-ready ZIP bundle,
SCAO overlays (MC 20, MC 12) *only if data is complete*, Judge profile seed.

──────────────────────────────────────────────────────────────────────────────
SETUP (PowerShell)
python -m venv .venv
.\.venv\Scripts\activate
pip install --upgrade pip
pip install pypdf python-docx

# Place this file next to golden_litigator_os.py and the DB it creates:
#   .\golden_litigator.db
#   .\LegalResults\...

RUN
python .\litigator_upgrade_suite.py

OUTPUTS
- LegalResults\Binder\Binder_Index.docx
- LegalResults\Binder\Binder_Combined.pdf (if any PDFs found)
- LegalResults\Bundles\MiFile_Package.zip
- LegalResults\Judges\judge_profiles.json (seeded if judge exists)
- LegalResults\Forms\MC20_FeeWaiver.docx (only if required data present)
- LegalResults\Forms\MC12_ProofOfService.docx (only if required data present)

POLICY
- No placeholders: forms are emitted ONLY when required fields are present.
- No deletion of source files.
"""

import os, json, sqlite3, zipfile
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Inches
from pypdf import PdfReader, PdfWriter

DB_PATH = "golden_litigator.db"
RESULTS = Path("LegalResults")
BINDER_DIR = RESULTS / "Binder"
FORMS_DIR = RESULTS / "Forms"
BUNDLES_DIR = RESULTS / "Bundles"
JUDGES_DIR = RESULTS / "Judges"

EXHIBIT_PDF_OUT = BINDER_DIR / "Binder_Combined.pdf"
EXHIBIT_INDEX_DOCX = BINDER_DIR / "Binder_Index.docx"
ZIP_OUT = BUNDLES_DIR / "MiFile_Package.zip"
JUDGE_JSON = JUDGES_DIR / "judge_profiles.json"


def ensure_dirs():
    for d in [RESULTS, BINDER_DIR, FORMS_DIR, BUNDLES_DIR, JUDGES_DIR]:
        d.mkdir(parents=True, exist_ok=True)


def db_conn():
    return sqlite3.connect(DB_PATH)


def load_case_meta() -> Optional[Dict[str, Any]]:
    if not Path(DB_PATH).exists():
        return None
    conn = db_conn()
    cur = conn.cursor()
    cur.execute(
        """SELECT court_name, case_number, caption_plaintiff, caption_defendant, judge, jurisdiction, division
                   FROM case_meta ORDER BY id DESC LIMIT 1"""
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        return None
    return {
        "court_name": row[0] or "",
        "case_number": row[1] or "",
        "caption_plaintiff": row[2] or "",
        "caption_defendant": row[3] or "",
        "judge": row[4] or "",
        "jurisdiction": row[5] or "",
        "division": row[6] or "",
    }


def fetch_top_evidence(limit=300) -> List[Dict[str, Any]]:
    if not Path(DB_PATH).exists():
        return []
    conn = db_conn()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT filename, filepath, ext, parties_json, claims_json, statutes_json, court_rules_json, exhibit_label
        FROM evidence ORDER BY relevance_score DESC, id DESC LIMIT ?
    """,
        (limit,),
    )
    rows = cur.fetchall()
    conn.close()
    out = []
    for r in rows:
        out.append(
            {
                "filename": r[0],
                "filepath": r[1],
                "ext": (r[2] or "").lower(),
                "parties": json.loads(r[3] or "[]"),
                "claims": json.loads(r[4] or "[]"),
                "statutes": json.loads(r[5] or "[]"),
                "rules": json.loads(r[6] or "[]"),
                "label": r[7] or "",
            }
        )
    return out


def next_label(idx: int) -> str:
    s = ""
    idx0 = idx
    while True:
        idx0, rem = divmod(idx0, 26)
        s = chr(65 + rem) + s
        if idx0 == 0:
            break
        idx0 -= 1
    return s


def build_binder_index(evidence: List[Dict[str, Any]]):
    doc = Document()
    doc.add_heading("MASTER EXHIBIT INDEX", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Exhibit"
    hdr[1].text = "File"
    hdr[2].text = "Path"
    hdr[3].text = "Claims/Statutes (Summary)"
    hdr[4].text = "Notes"

    ex_count = 0
    for i, rec in enumerate(evidence):
        label = rec["label"].strip() if rec["label"] else next_label(ex_count)
        ex_count += 1

        cl_summary = []
        if rec["claims"]:
            cl_summary.append(
                "Claims:"
                + ", ".join(
                    [
                        c if isinstance(c, str) else c.get("name", "")
                        for c in rec["claims"]
                    ][:3]
                )
            )
        if rec["statutes"]:
            cl_summary.append(
                "Stat:"
                + ", ".join(
                    [
                        s if isinstance(s, str) else s.get("cite", "")
                        for s in rec["statutes"]
                    ][:3]
                )
            )
        row = table.add_row().cells
        row[0].text = f"Exhibit {label}"
        row[1].text = rec["filename"]
        row[2].text = rec["filepath"]
        row[3].text = " | ".join(cl_summary) if cl_summary else ""
        row[4].text = ""
    doc.save(str(EXHIBIT_INDEX_DOCX))


def merge_pdf_binder(evidence: List[Dict[str, Any]]):
    writer = PdfWriter()
    any_pdf = False
    for rec in evidence:
        p = Path(rec["filepath"])
        if p.suffix.lower() == ".pdf" and p.exists():
            try:
                reader = PdfReader(str(p))
                for page in reader.pages:
                    writer.add_page(page)
                any_pdf = True
            except Exception:
                continue
    if any_pdf:
        with open(EXHIBIT_PDF_OUT, "wb") as f:
            writer.write(f)


def seed_judge_profile(meta: Dict[str, Any]):
    if not meta or not meta.get("judge"):
        return
    JUDGES_DIR.mkdir(parents=True, exist_ok=True)
    data = []
    if JUDGE_JSON.exists():
        try:
            data = json.loads(JUDGE_JSON.read_text(encoding="utf-8", errors="ignore"))
        except Exception:
            data = []
    names = {j.get("name") for j in data}
    if meta["judge"] not in names:
        data.append(
            {
                "name": meta["judge"],
                "court": meta.get("court_name", ""),
                "division": meta.get("division", ""),
                "notes": "Seeded from case_meta; expand with tendencies, timing, sanctions history.",
            }
        )
        JUDGE_JSON.write_text(
            json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
        )


def generate_mc20_fee_waiver(meta: Dict[str, Any]):
    needed = [
        meta.get("court_name", ""),
        (meta.get("caption_defendant") or meta.get("caption_plaintiff")),
        meta.get("case_number"),
    ]
    if not all(needed):
        return
    name = (meta.get("caption_defendant") or meta.get("caption_plaintiff")).strip()
    doc = Document()
    doc.add_heading("MC 20 — REQUEST AND ORDER FOR FEE WAIVER (Draft)", level=1)
    doc.add_paragraph(f"Court: {meta.get('court_name','')}")
    doc.add_paragraph(f"Case No.: {meta.get('case_number','')}")
    doc.add_paragraph(f"Applicant: {name}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "This draft is generated for completion. Attach financial affidavit and supporting proofs per MCR 2.002."
    )
    FORMS_DIR.mkdir(parents=True, exist_ok=True)
    out = FORMS_DIR / "MC20_FeeWaiver.docx"
    doc.save(str(out))


def generate_mc12_proof_of_service(
    meta: Dict[str, Any], served_doc_title: Optional[str] = None
):
    needed = [
        meta.get("case_number"),
        (meta.get("caption_defendant") or meta.get("caption_plaintiff")),
    ]
    if not all(needed):
        return
    doc = Document()
    doc.add_heading("MC 12 — PROOF OF SERVICE (Draft)", level=1)
    doc.add_paragraph(f"Case No.: {meta.get('case_number','')}")
    if served_doc_title:
        doc.add_paragraph(f"Document Served: {served_doc_title}")
    doc.add_paragraph(
        "Complete service method, date, and server details per MCR 2.107 / 2.105 as applicable."
    )
    out = FORMS_DIR / "MC12_ProofOfService.docx"
    doc.save(str(out))


def build_mifile_zip():
    files_to_zip = []
    for p in [
        EXHIBIT_INDEX_DOCX,
        EXHIBIT_PDF_OUT,
        RESULTS / "Narratives",
        RESULTS / "Motions",
        RESULTS / "EvidenceLedger.jsonl",
    ]:
        if p.is_file():
            files_to_zip.append(p)
        elif p.is_dir():
            for f in p.rglob("*"):
                if f.is_file():
                    files_to_zip.append(f)
    if not files_to_zip:
        return
    with zipfile.ZipFile(ZIP_OUT, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for f in files_to_zip:
            try:
                arc = f.relative_to(RESULTS) if f.is_relative_to(RESULTS) else f.name
            except Exception:
                arc = f.name
            z.write(f, arcname=str(arc))


def main():
    ensure_dirs()
    meta = load_case_meta()

    ev = fetch_top_evidence(limit=400)
    if ev:
        build_binder_index(ev)
        merge_pdf_binder(ev)

    if meta:
        seed_judge_profile(meta)
        generate_mc20_fee_waiver(meta)
        generate_mc12_proof_of_service(meta)

    build_mifile_zip()

    print("✅ Upgrade suite complete:")
    print(
        f" - Binder Index: {EXHIBIT_INDEX_DOCX.exists() and EXHIBIT_INDEX_DOCX or 'skipped'}"
    )
    print(
        f" - Binder PDF:   {EXHIBIT_PDF_OUT.exists() and EXHIBIT_PDF_OUT or 'none found'}"
    )
    print(f" - MiFile ZIP:   {ZIP_OUT.exists() and ZIP_OUT or 'skipped'}")
    print(f" - Judge JSON:   {JUDGE_JSON.exists() and JUDGE_JSON or 'skipped'}")
    print(f" - Forms (MC20/MC12): in {FORMS_DIR} (emitted only if fields complete)")


if __name__ == "__main__":
    main()
