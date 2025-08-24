"""Draft overlay generators for common SCAO forms."""

from __future__ import annotations
from pathlib import Path
from typing import Dict, Optional

from docx import Document

from utils import db_conn

FORMS_DIR = Path("LegalResults/Forms")
FORMS_DIR.mkdir(parents=True, exist_ok=True)


def _load_case() -> Optional[Dict[str, str]]:
    conn = db_conn("golden_litigator.db")
    cur = conn.cursor()
    cur.execute(
        """SELECT court_name, case_number, caption_plaintiff, caption_defendant
        FROM case_meta ORDER BY id DESC LIMIT 1"""
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        return None
    return {
        "court_name": row[0] or "",
        "case_number": row[1] or "",
        "pl": row[2] or "",
        "df": row[3] or "",
    }


def mc20_fee_waiver() -> None:
    meta = _load_case()
    if (
        not meta
        or not meta["court_name"]
        or not meta["case_number"]
        or not (meta["pl"] or meta["df"])
    ):
        return
    doc = Document()
    doc.add_heading("MC 20 (Draft Overlay) – Request and Order for Fee Waiver", level=1)
    doc.add_paragraph(
        f"Court: {meta['court_name']}  |  Case No.: {meta['case_number']}"
    )
    doc.add_paragraph(f"Applicant: {meta['df'] or meta['pl']}")
    doc.add_paragraph("Attach financial affidavit and proofs per MCR 2.002.")
    doc.save(str(FORMS_DIR / "MC20_FeeWaiver_Draft.docx"))


def mc12_proof_of_service(served_title: str = "") -> None:
    meta = _load_case()
    if not meta or not meta["case_number"] or not (meta["pl"] or meta["df"]):
        return
    doc = Document()
    doc.add_heading("MC 12 (Draft Overlay) – Proof of Service", level=1)
    doc.add_paragraph(f"Case No.: {meta['case_number']}")
    if served_title:
        doc.add_paragraph(f"Document Served: {served_title}")
    doc.add_paragraph("Complete service details per MCR 2.105/2.107.")
    doc.save(str(FORMS_DIR / "MC12_ProofOfService_Draft.docx"))


if __name__ == "__main__":
    mc20_fee_waiver()
    mc12_proof_of_service("Motion/Brief")
