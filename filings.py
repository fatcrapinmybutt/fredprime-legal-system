"""Motion document builder."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from court_engine import requirements_met
from agents import call_llm


def make_motion_docx(
    results_dir: Path,
    motion_type: str,
    case_meta: Dict[str, Any],
    materials: Dict[str, Any],
) -> Optional[Path]:
    """Generate motion DOCX when requirements satisfied."""
    if not requirements_met(motion_type, materials, case_meta):
        return None

    motions_dir = results_dir / "Motions"
    motions_dir.mkdir(parents=True, exist_ok=True)
    out_path = (
        motions_dir
        / f"{motion_type.replace(' ', '_')}_{case_meta.get('case_number', 'NOCASE')}.docx"
    )

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_paragraph()
    run = title.add_run(
        case_meta.get(
            "court_name", "IN THE CIRCUIT COURT FOR MUSKEGON COUNTY, MICHIGAN"
        )
    )
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    cap = doc.add_paragraph()
    cap.add_run(f"{case_meta.get('caption_plaintiff', '')}, Plaintiff,\n").bold = True
    cap.add_run("v.\n")
    cap.add_run(f"{case_meta.get('caption_defendant', '')}, Defendant.\n").bold = True
    if case_meta.get("case_number"):
        doc.add_paragraph(f"Case No.: {case_meta['case_number']}")
    if case_meta.get("judge"):
        doc.add_paragraph(f"Hon.: {case_meta['judge']}")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run(motion_type.upper())
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    payload = {
        "mtype": motion_type,
        "case_meta": case_meta,
        "materials": materials.get("materials", [])[:60],
        "style_requirements": {
            "jurisdiction": "Michigan/W.D. Mich as applicable",
            "format": "Court-ready, numbered, accurate cites only, no placeholders",
            "relief": "Specific",
        },
    }
    prompt = (
        "Draft a complete motion as strict JSON: "
        '{"sections":[{"heading":"...","paragraphs":["..."]}], '
        '"relief":["..."], "proposed_order":["..."]} from: \n'
        + json.dumps(payload, ensure_ascii=False)
    )
    body_json = call_llm(prompt)
    try:
        body = json.loads(body_json)
    except Exception:
        return None

    for sec in body.get("sections", []):
        if sec.get("heading"):
            h = doc.add_paragraph()
            hr = h.add_run(sec["heading"])
            hr.bold = True
        for para in sec.get("paragraphs", []):
            par = doc.add_paragraph(para)
            for run in par.runs:
                run.font.size = Pt(12)

    if body.get("relief"):
        doc.add_paragraph("")
        rr = doc.add_paragraph()
        rr.add_run("RELIEF REQUESTED").bold = True
        for item in body["relief"]:
            doc.add_paragraph(f"• {item}")

    if body.get("proposed_order"):
        doc.add_paragraph("")
        po = doc.add_paragraph()
        po.add_run("PROPOSED ORDER (ATTACHED)").bold = True
        for line in body["proposed_order"]:
            doc.add_paragraph(line)

    doc.save(str(out_path))
    return out_path
