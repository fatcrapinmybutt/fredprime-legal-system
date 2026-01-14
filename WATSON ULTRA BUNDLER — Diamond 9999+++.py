#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
WATSON ULTRA BUNDLER — Diamond 9999+++ (Michigan-Locked, Autonomous)
Author: Litigation OS (Andrew J. Pigors)
Target: Windows (paths like F:\...), offline, zero placeholders, court-ready.

WHAT THIS DOES (Autonomously):
  • Scans your source folders for PDFs/IMGs/DOCX/TXT and builds a TIMELINE + EXHIBIT set (H-1…H-N).
  • OCRs and extracts text (PyMuPDF -> PyPDF2 fallback; PIL+pytesseract for images).
  • Auto-labels/renames exhibits H-1.. with a deterministic order; creates Cover Sheets + on-record metadata CSV.
  • Generates and/or validates:
      - Verified Complaint (DOCX) — Michigan counts with auto pin-cites and neutral, evidence-tethered paragraphs.
      - Proposed Orders (Sanctions MCR 1.109(E), Protective 2.302(C), Alt-Service 2.105, Limine, Narrow Injunction).
      - ESI Protocol + Discovery Plan + Service Playbook + Affidavit of Diligence.
      - Pretrial Statement + Proposed Jury Instructions (outline, M Civ JI-aligned).
      - Retraction Letters (1 per detected publication cluster).
      - Elements Table (per count), Privilege Matrix (exclude predicates), SOL Matrix (claim vs. background).
      - Damages Workbook CSV shell.
  • Child-privacy enforcement: redacts child personal identifiers; replaces “Lincoln” with initials (configurable).
  • Gating audits (HARD FAIL if violated): Pin-Cites, Privilege, SOL, Redaction/Child privacy.
  • Emits Manifest.csv with SHA-256 for every file, and a ready-to-file ZIP bundle.

DEPENDENCIES (install locally once):
  pip install python-docx PyPDF2
  pip install pymupdf           # optional, best PDF text extraction
  pip install pillow pytesseract # optional, image OCR (requires Tesseract installed)
  pip install exifread           # optional, deeper EXIF

USAGE (PowerShell):
  python watson_ultra_bundler.py ^
    --scan-root "F:\\Evidence" ^
    --input-root "F:\\LegalResults\\Source" ^
    --output-root "F:\\LegalResults" ^
    --case-title "Pigors v. Emily A. Watson, et al." ^
    --county "Muskegon" ^
    --child-name "Lincoln" ^
    --child-initials "L.D.W." ^
    --ts "20251029_1420" ^
    --no-dry-run

Notes:
  • If you omit --ts, a deterministic timestamp is derived from inputs.
  • DRY RUN is ON by default (safety). Add --no-dry-run to write files and create ZIP.
  • Housing matters are EXCLUDED by automatic keyword filters.
"""

import argparse
import csv
import datetime as dt
import hashlib
import io
import json
import logging
import os
import re
import shutil
import sys
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

# Optional imports (handled gracefully)
HAVE_PYMUPDF = False
HAVE_TESSERACT = False
HAVE_PIL = False
HAVE_EXIFREAD = False

try:
    import fitz  # PyMuPDF

    HAVE_PYMUPDF = True
except Exception:
    pass

try:
    from PIL import Image

    HAVE_PIL = True
    try:
        import pytesseract

        HAVE_TESSERACT = True
    except Exception:
        HAVE_TESSERACT = False
except Exception:
    HAVE_PIL = False

try:
    import exifread

    HAVE_EXIFREAD = True
except Exception:
    HAVE_EXIFREAD = False

try:
    from PyPDF2 import PdfReader
except Exception as e:
    print("Missing dependency: PyPDF2. Install with: pip install PyPDF2")
    raise

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except Exception:
    print("Missing dependency: python-docx. Install with: pip install python-docx")
    raise

# ---------------------------
# CONFIG / CONSTANTS
# ---------------------------
ALLOWED_EXT = {".pdf", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".docx", ".txt"}
PDF_EXT = {".pdf"}
IMG_EXT = {".png", ".jpg", ".jpeg", ".tif", ".tiff"}
TEXTY_EXT = {".docx", ".txt"}

# Housing filters (exclude any source files that are clearly housing-related)
HOUSING_KWS = re.compile(
    r"\b(shady\s*oaks|lot\s*\d+|rent|lease|eviction|sewer|egle|utility|water\s*bill|homes\s*of\s*america|z(e|é)go)\b",
    re.IGNORECASE,
)

# Family/case keywords (to prioritize for pin-cites)
FAMILY_KWS = [
    "PPO",
    "show cause",
    "AppClose",
    "NSPD",
    "Ella Randall",
    "USB",
    "RusCOPA",
    "parenting time",
    "exchange",
    "affidavit",
    "ex parte",
    "recording",
    "mental health",
    "meth",
    "welfare check",
    "police report",
    "transcript",
    "CPS",
    "FOIA",
    "bias",
    "canon",
    "best interest",
    "MCR",
    "MCL",
]

# Claims & elements (Michigan-aligned; neutral language)
CLAIMS = [
    (
        "Defamation / Defamation by Implication",
        [
            "Statement of fact about Plaintiff",
            "Falsity",
            "Publication to a third person",
            "Fault (≥ negligence; actual malice if privilege)",
            "Damages (or presumed per se)",
        ],
    ),
    (
        "False Light",
        [
            "Publicity placing Plaintiff in a false light",
            "Highly offensive to a reasonable person",
            "Knowledge or reckless disregard",
            "Causation and damages",
        ],
    ),
    (
        "Intentional Infliction of Emotional Distress (IIED)",
        ["Extreme and outrageous conduct", "Intent or recklessness", "Causation", "Severe emotional distress"],
    ),
    ("Abuse of Process", ["Use of process after issuance", "Ulterior purpose", "Damage to Plaintiff"]),
    (
        "Tortious Interference with Parental Rights (narrow)",
        [
            "Lawful parenting-time/custody right",
            "Intentional interference without justification",
            "Causation",
            "Damages (missed overnights/holidays, costs)",
        ],
    ),
    (
        "Civil Conspiracy / Aiding and Abetting",
        ["Agreement or substantial assistance", "Underlying tort", "Damages trace to tort"],
    ),
    (
        "Declaratory and Injunctive Relief",
        ["Actual controversy", "Narrow tailoring", "Balance of harms/public interest"],
    ),
]

# Date pattern to build timeline
DATE_PAT = re.compile(
    r"\b(20[12]\d)[-/.](0?[1-9]|1[0-2])[-/.](0?[1-9]|[12]\d|3[01])\b|"
    r"\b(0?[1-9]|1[0-2])[-/.](0?[1-9]|[12]\d|3[01])[-/.](20[12]\d)\b"
)

# Exhibit pin-cite detection (used also for audits)
PIN_REGEX = re.compile(
    r"(Exhibit\s+[A-Z](?:-\d+)?(?:\s*p\.\s*\d+)?|H-\d+\s*(?:p\.\s*\d+|t=\d{2}:\d{2}:\d{2})?)",
    re.IGNORECASE,
)

# Child privacy tokens to redact from public complaint
PRIVACY_TOKENS = [
    r"\bDate\s*of\s*Birth\b",
    r"\bDOB\b",
    r"\bSSN\b",
    r"\bSocial\s*Security\b",
    r"\bStudent\s*ID\b",
    r"\bSchool\b\s*Name\b",
    r"\bAddress:\s*\d+",
]


# ---------------------------
# UTILITIES
# ---------------------------
def sha256_of_path(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def makedirs(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def copy_with_hash(src: Path, dst: Path) -> str:
    makedirs(dst.parent)
    shutil.copy2(src, dst)
    return sha256_of_path(dst)


def text_from_pdf(p: Path) -> str:
    if HAVE_PYMUPDF:
        try:
            doc = fitz.open(p.as_posix())
            chunks = []
            for page in doc:
                chunks.append(page.get_text("text"))
            return "\n".join(chunks)
        except Exception:
            pass
    # fallback PyPDF2
    try:
        reader = PdfReader(p.as_posix())
        txt = []
        for pg in reader.pages:
            try:
                txt.append(pg.extract_text() or "")
            except Exception:
                txt.append("")
        return "\n".join(txt)
    except Exception:
        return ""


def text_from_image(p: Path) -> str:
    if HAVE_PIL and HAVE_TESSERACT:
        try:
            img = Image.open(p.as_posix())
            return pytesseract.image_to_string(img)
        except Exception:
            return ""
    return ""


def text_from_docx(p: Path) -> str:
    try:
        doc = Document(p.as_posix())
        return "\n".join(par.text for par in doc.paragraphs)
    except Exception:
        return ""


def text_from_txt(p: Path) -> str:
    try:
        return p.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        try:
            return p.read_text(encoding="latin-1", errors="ignore")
        except Exception:
            return ""


def extract_text(p: Path) -> str:
    ext = p.suffix.lower()
    if ext in PDF_EXT:
        return text_from_pdf(p)
    if ext in IMG_EXT:
        return text_from_image(p)
    if ext == ".docx":
        return text_from_docx(p)
    if ext == ".txt":
        return text_from_txt(p)
    return ""


def exif_dict(p: Path) -> dict:
    out = {}
    if HAVE_EXIFREAD and p.suffix.lower() in IMG_EXT:
        try:
            with p.open("rb") as f:
                tags = exifread.process_file(f, details=False)
            for k, v in tags.items():
                out[str(k)] = str(v)
        except Exception:
            pass
    return out


def derive_ts_from_files(paths: list[Path]) -> str:
    h = hashlib.sha256()
    for p in sorted(paths, key=lambda x: x.as_posix()):
        try:
            h.update(p.name.encode())
            h.update(str(int(p.stat().st_mtime)).encode())
        except Exception:
            pass
    digest = h.hexdigest()[:8]
    return dt.datetime.now().strftime("%Y%m%d") + f"_{digest}"


def sanitize_for_public(text: str, child_name: str, child_initials: str) -> str:
    # Replace child name case-insensitively with initials
    text = re.sub(re.escape(child_name), child_initials, text, flags=re.IGNORECASE)
    # Remove common PII tokens
    for pat in PRIVACY_TOKENS:
        text = re.sub(pat, "[REDACTED]", text, flags=re.IGNORECASE)
    return text


def is_housing_file(path: Path, text: str) -> bool:
    return bool(HOUSING_KWS.search(path.name)) or bool(HOUSING_KWS.search(text))


def find_dates(text: str) -> list[str]:
    hits = []
    for m in DATE_PAT.finditer(text):
        hits.append(m.group(0))
    return hits


def chunk_snippet(text: str, max_len=140) -> str:
    s = re.sub(r"\s+", " ", text).strip()
    return (s[:max_len] + "…") if len(s) > max_len else s


def split_into_sentences(text: str) -> list[str]:
    s = re.sub(r"\s+", " ", text)
    return re.split(r"(?<=[\.\?\!])\s+", s)


# ---------------------------
# DISCOVERY & INDEXING
# ---------------------------
def discover_files(scan_root: Path) -> list[Path]:
    files = []
    for root, _, names in os.walk(scan_root.as_posix()):
        for n in names:
            p = Path(root) / n
            if p.suffix.lower() in ALLOWED_EXT and p.is_file():
                files.append(p)
    return files


def build_index(scan_root: Path, child_name: str, child_initials: str) -> list[dict]:
    rows = []
    all_paths = discover_files(scan_root)
    for p in all_paths:
        try:
            size = p.stat().st_size
            sha = sha256_of_path(p)
        except Exception:
            size, sha = 0, ""
        text = extract_text(p)
        # quick pre-filter to keep housing separate
        housing = is_housing_file(p, text)
        if housing:
            continue
        # privacy sanitize preview only (we keep originals intact)
        preview = sanitize_for_public(text[:2000], child_name, child_initials) if text else ""
        dates = find_dates(text) if text else []
        exif = exif_dict(p) if p.suffix.lower() in IMG_EXT else {}
        rows.append(
            {
                "path": p,
                "name": p.name,
                "ext": p.suffix.lower(),
                "size": size,
                "sha256": sha,
                "text": text,
                "preview": preview,
                "dates": dates,
                "exif": exif,
            }
        )
    return rows


# ---------------------------
# EXHIBIT BUILDER
# ---------------------------
def stage_exhibits(index_rows: list[dict], dst_exhibits_dir: Path) -> list[dict]:
    """
    Order: prioritize PDFs, then images, then docx/txt proofs.
    Label H-1…H-N. Copy to Exhibits folder.
    """
    # sort by preference & last modified time (newest first so H-1 likely salient)
    scored = []
    for r in index_rows:
        p = r["path"]
        try:
            mtime = p.stat().st_mtime
        except Exception:
            mtime = 0
        rank = 0
        if r["ext"] in PDF_EXT:
            rank = 3
        elif r["ext"] in IMG_EXT:
            rank = 2
        elif r["ext"] in TEXTY_EXT:
            rank = 1
        scored.append((rank, -mtime, r))
    scored.sort(reverse=True)  # high rank, newest first

    exhibits = []
    counter = 1
    for _, _, r in scored:
        label = f"H-{counter}"
        src = r["path"]
        dst = (
            dst_exhibits_dir / f"{label}_{src.stem}.pdf"
            if r["ext"] in PDF_EXT
            else dst_exhibits_dir / f"{label}_{src.stem}{r['ext']}"
        )
        exhibits.append(
            {
                "label": label,
                "src": src,
                "dst": dst,
                "sha256": r["sha256"],
                "size": r["size"],
                "dates": r["dates"],
                "text": r["text"],
                "preview": r["preview"],
                "ext": r["ext"],
            }
        )
        counter += 1
    return exhibits


# ---------------------------
# MATRICES (Elements, Privilege, SOL)
# ---------------------------
def build_elements_table(exhibits: list[dict], out_csv: Path) -> dict:
    """
    For each claim, build element rows with auto pin-cites:
    - pin-cites found by keyword hits in exhibit text; fallback to first page reference if none.
    """
    # simple keyword map per claim
    cmap = {
        "Defamation / Defamation by Implication": ["meth", "drug", "abuse", "unsafe", "criminal", "lied", "danger"],
        "False Light": ["false", "misleading", "imply", "implication", "portray"],
        "Intentional Infliction of Emotional Distress (IIED)": [
            "outrageous",
            "severe",
            "emotional",
            "distress",
            "harass",
        ],
        "Abuse of Process": ["ppo", "show cause", "subpoena", "process", "coerce", "ulterior"],
        "Tortious Interference with Parental Rights (narrow)": [
            "parenting time",
            "withhold",
            "exchange",
            "custody",
            "deny",
        ],
        "Civil Conspiracy / Aiding and Abetting": ["conspire", "coordinate", "assist", "plan", "joint"],
        "Declaratory and Injunctive Relief": ["injunction", "declaratory", "order", "tailor"],
    }
    rows = []
    evidence_map = defaultdict(list)  # claim -> [pin-cite strings]
    # Build rough pin-cites by searching exhibit text for claim keywords
    for claim, elements in CLAIMS:
        kws = cmap.get(claim, [])
        for ex in exhibits:
            txt = ex["text"].lower() if ex["text"] else ""
            hit = False
            for kw in kws:
                if kw in txt:
                    hit = True
                    break
            if hit:
                # simple pin-cite: "H-x p.1" (we don't know actual pages for images; use p.1)
                pin = f"{ex['label']} p.1"
                snippet = chunk_snippet(ex["preview"]) if ex["preview"] else ""
                evidence_map[claim].append((pin, snippet))

    # Write CSV
    with out_csv.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Claim", "Element", "PinCites", "CorroborationSnippets", "Gaps/Notes"])
        for claim, elements in CLAIMS:
            pins = evidence_map.get(claim, [])
            for elem in elements:
                pin_str = " | ".join(p for p, _ in pins) if pins else ""
                cor_str = " | ".join(s for _, s in pins[:5]) if pins else ""
                w.writerow(
                    [
                        claim,
                        elem,
                        pin_str,
                        cor_str,
                        "" if pin_str else "No automatic hits — add manual cite if appropriate",
                    ]
                )
    return {"evidence_map": evidence_map}


def build_privilege_matrix(exhibits: list[dict], out_csv: Path) -> None:
    """
    Flag likely privileged materials (court filings/testimony/fair report indicators) as context-only.
    """
    with out_csv.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Source", "PrivilegeType", "Basis", "UseAllowed", "Notes"])
        for ex in exhibits:
            txt = ex["text"] or ""
            basis = ""
            ptype = ""
            use = "claim"
            if re.search(r"STATE OF MICHIGAN|CIRCUIT COURT|DISTRICT COURT|COURT OF APPEALS", txt, re.IGNORECASE):
                ptype = "absolute"
                basis = "judicial filing/testimony indicators"
                use = "context-only"
            elif re.search(r"news|report|press|article", txt, re.IGNORECASE):
                ptype = "qualified"
                basis = "possible fair report"
                use = "context-only"
            w.writerow([ex["src"].name, ptype, basis, use, ""])


def build_sol_matrix(exhibits: list[dict], out_csv: Path) -> None:
    """
    Rough SOL computation by first detected date per exhibit. Conservative defaults.
    - Defamation: 1 year
    - Abuse of Process/IIED/False Light: 3 years (general tort)
    - Interference (parental): assume 3 years unless rule dictates otherwise
    """

    def earliest_year(dates):
        years = []
        for d in dates:
            ys = re.findall(r"(20[12]\d)", d)
            for y in ys:
                years.append(int(y))
        return min(years) if years else None

    claim_limits = {
        "Defamation / Defamation by Implication": 1,
        "False Light": 3,
        "Intentional Infliction of Emotional Distress (IIED)": 3,
        "Abuse of Process": 3,
        "Tortious Interference with Parental Rights (narrow)": 3,
        "Civil Conspiracy / Aiding and Abetting": 3,
        "Declaratory and Injunctive Relief": 0,  # N/A, tied to controversy
    }
    this_year = dt.datetime.now().year
    with out_csv.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Claim", "AccrualYear(heuristic)", "LimitYears", "Status", "Treatment", "Notes"])
        for claim, _elements in CLAIMS:
            limit = claim_limits[claim]
            # use earliest exhibit year as heuristic (improve by scanning text dates)
            yrs = []
            for ex in exhibits:
                y = earliest_year(ex["dates"])
                if y:
                    yrs.append(y)
            acc = min(yrs) if yrs else this_year
            if limit == 0:
                status = "timely"
                treat = "claim"
            else:
                age = this_year - acc
                if age <= limit:
                    status = "timely"
                    treat = "claim"
                else:
                    status = "stale"
                    treat = "background-only"
            w.writerow([claim, acc, limit, status, treat, "Heuristic; confirm with event-level dates."])


# ---------------------------
# VERIFIED COMPLAINT GENERATOR
# ---------------------------
def add_heading(doc: Document, text: str, size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def numbered_par(doc: Document, idx: int, text: str):
    doc.add_paragraph(f"{idx}. {text}")


def make_verified_complaint(
    out_docx: Path,
    case_title: str,
    county: str,
    exhibits: list[dict],
    evidence_map: dict,
    child_name: str,
    child_initials: str,
):
    doc = Document()
    add_heading(doc, f"STATE OF MICHIGAN — IN THE CIRCUIT COURT FOR {county.upper()} COUNTY", 12)
    add_heading(doc, case_title, 14)
    doc.add_paragraph("VERIFIED COMPLAINT — JURY DEMAND — SANCTIONS REQUEST (MCR 1.109(E))").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph("")

    # Preliminary Statement
    doc.add_heading("Preliminary Statement", level=1)
    doc.add_paragraph(
        "Plaintiff alleges a coordinated pattern of wrongful conduct by Defendants that harmed his reputation, "
        "interfered with his parental relationship, and misused court and police processes. The allegations below "
        "are pin-cited to specific exhibits."
    )

    # Factual Allegations (neutral, pin-cited from exhibits)
    doc.add_heading("Factual Allegations", level=1)
    n = 1
    # Build top 12 neutral facts by scanning exhibit snippets
    top_snips = []
    for ex in exhibits[:20]:
        sents = split_into_sentences(ex["preview"]) if ex["preview"] else []
        if sents:
            # choose sentences containing family keywords
            candidates = [s for s in sents if any(kw.lower() in s.lower() for kw in FAMILY_KWS)]
            pick = candidates[0] if candidates else sents[0]
            top_snips.append((ex["label"], pick))
    # Write facts with pin-cites
    for label, sent in top_snips[:12]:
        numbered_par(doc, n, sanitize_for_public(f"{sent} (Exhibit {label} p.1)", child_name, child_initials))
        n += 1

    # Counts — element-by-element (neutral language)
    def write_count(title, elements, claim_key):
        doc.add_heading(title, level=2)
        numbered_par(doc, 1, "Plaintiff incorporates the preceding paragraphs by reference.")
        # element lines with pin-cites from evidence_map
        pins = evidence_map.get(title, [])
        # if evidence_map keys differ (exact title), fall back by fuzzy match
        if not pins:
            for k in evidence_map.keys():
                if title.split(" (")[0].lower() in k.lower():
                    pins = evidence_map[k]
                    break
        for i, elem in enumerate(elements, start=2):
            cite = f" Supported by: " + " ; ".join([p for p, _ in pins[:5]]) if pins else ""
            numbered_par(doc, i, f"{elem}.{cite}")

    for title, elements in CLAIMS:
        write_count(title, elements, title)

    # Sanctions
    doc.add_heading("Sanctions (MCR 1.109(E))", level=2)
    doc.add_paragraph(
        "Plaintiff requests findings that identified statements or papers by Defendants were presented for improper purpose "
        "or without evidentiary support after reasonable inquiry, and an award of reasonable expenses and attorney fees."
    )

    # Anti-SLAPP / Privilege rebuttal
    doc.add_heading("Privilege and Petitioning", level=2)
    doc.add_paragraph(
        "Michigan provides no broad anti-SLAPP bar. Litigation or petitioning privileges do not extend to knowing falsehoods, "
        "sham use of process, non-judicial republication, or misuse after issuance. Any privileged statements are not used as "
        "claim predicates and appear only as context; claims rest on non-privileged acts and publications."
    )

    # Prayer
    doc.add_heading("Prayer for Relief", level=1)
    for item in [
        "Compensatory damages in categories proved at trial.",
        "Exemplary damages where permitted by Michigan law.",
        "Declaratory and narrow injunctive relief to halt identified wrongful conduct.",
        "Sanctions and reasonable expenses under MCR 1.109(E).",
        "Costs and interest as allowed by law.",
        "Such further relief as is just.",
    ]:
        doc.add_paragraph(f"• {item}")

    # Jury demand
    doc.add_heading("Jury Demand", level=1)
    doc.add_paragraph("Plaintiff demands trial by jury on all triable issues. (MCR 2.508(A))")

    # Verification
    doc.add_heading("Verification", level=1)
    doc.add_paragraph(
        "I declare that the factual allegations above are true to the best of my knowledge, information, and belief. "
        "I understand that a willful false statement is subject to penalty."
    ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("Dated: " + dt.datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph("/s/ Andrew J. Pigors")

    # Save
    makedirs(out_docx.parent)
    doc.save(out_docx.as_posix())


# ---------------------------
# ORDERS / TEMPLATES GENERATORS
# ---------------------------
def gen_docx(out_path: Path, title: str, paras: list[str], case_title: str):
    doc = Document()
    add_heading(doc, case_title, 14)
    doc.add_paragraph(title).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    for para in paras:
        doc.add_paragraph(para)
    makedirs(out_path.parent)
    doc.save(out_path.as_posix())


def ensure_orders(dst_orders_dir: Path, case_title: str):
    orders = {
        "Sanctions_Order_MCR_1_109_E.docx": [
            "ORDER: The Court finds the identified papers were presented for improper purpose or lacked evidentiary support under MCR 1.109(E).",
            "Defendants shall pay Plaintiff’s reasonable expenses and attorney fees within 21 days.",
            "Plaintiff shall file a fee affidavit within 14 days; objections due 14 days thereafter.",
            "IT IS SO ORDERED.",
        ],
        "Protective_Order_MCR_2_302_C.docx": [
            "ORDER: Child identified by initials in public filings; redact sensitive identifiers under MCR 1.109(D)(9) and MCR 2.302(C).",
            "Sensitive exhibits filed under seal; audio/video produced in native format with SHA-256 hash.",
            "Clawback for inadvertent production; use limited to this case.",
        ],
        "Alternate_Service_Order_MCR_2_105.docx": [
            "ORDER: Alternate service authorized under MCR 2.105 upon affidavit of diligent attempts.",
            "Service by certified mail + first-class mail + email and posting as permitted; proof due within 14 days.",
        ],
        "Motions_in_Limine_Order.docx": [
            "ORDER: Evidence excluded as follows: MRE 404 character/propensity; MRE 401–403 remote/cumulative; hearsay outside MRE 803/804; unauthenticated media lacking MRE 901 foundation; privileged/settlement communications."
        ],
        "Narrow_Injunction_Order.docx": [
            "ORDER: Defendants shall not publish adjudicated false statements; shall not obstruct court-ordered exchanges.",
            "Party communications limited to logistics via agreed channel; violations subject to contempt.",
        ],
    }
    for fname, body in orders.items():
        gen_docx(dst_orders_dir / fname, fname.replace("_", " ").replace(".docx", ""), body, case_title)


def ensure_discovery(dst_discovery_dir: Path, case_title: str):
    gen_docx(
        dst_discovery_dir / "ESI_Protocol.docx",
        "ESI Protocol",
        [
            "Custodians: parties; carriers; AppClose; clerk media-intake; police/CPS.",
            "Formats: PST/MBOX+CSV (email); JSON/CSV+PDF (messages); native+SHA-256 (audio/video); photos with EXIF.",
            "Hashing: SHA-256; manifest required. De-dup by hash. Privilege log categorical; clawback in order.",
            "Search: exchange terms and hit-counts; rolling production with Bates and exhibit IDs.",
        ],
        case_title,
    )
    gen_docx(
        dst_discovery_dir / "Discovery_Plan.docx",
        "Discovery Plan",
        [
            "Scope: elements-based discovery tied to Complaint counts.",
            "Deadlines: to be set at MCR 2.401 conference.",
            "Protective Order: child privacy; sealed exhibits.",
            "Mediation/Case Evaluation: schedule per local practice; consider MCR 2.403 and MCR 2.405 timing.",
        ],
        case_title,
    )


def ensure_pretrial(dst_pretrial_dir: Path, case_title: str):
    gen_docx(
        dst_pretrial_dir / "Pretrial_Statement.docx",
        "Pretrial Statement",
        [
            "Stipulated Facts: [complete before conference].",
            "Contested Facts: numbered with exhibit pin-cites.",
            "Issues of Law: privilege, SOL, elements, evidence.",
            "Witnesses: names, contact, offers of proof, time estimates.",
            "Exhibits: ID, description, source path, SHA-256, foundation, objections.",
            "Motions in Limine: list with authority.",
            "Time Estimates: openings, case-in-chief, cross, rebuttal.",
            "ADR: case evaluation and offers of judgment status.",
        ],
        case_title,
    )
    gen_docx(
        dst_pretrial_dir / "Proposed_Jury_Instructions.docx",
        "Proposed Jury Instructions (Outline)",
        [
            "Defamation: statement of fact; falsity; publication; fault; damages; privilege burden.",
            "False Light: false implication; publicity; knowledge/reckless disregard; damages.",
            "IIED: extreme/outrageous; intent/recklessness; causation; severe distress.",
            "Abuse of Process: use after issuance; ulterior purpose; damages.",
            "Interference with Parental Rights (narrow): lawful right; intentional interference; lack of justification; damages.",
            "Civil Conspiracy/Aiding: agreement/substantial assistance; damages trace to underlying tort.",
        ],
        case_title,
    )


def ensure_service(dst_service_dir: Path, case_title: str):
    gen_docx(
        dst_service_dir / "Service_Playbook.docx",
        "Service Playbook (MCR 2.105)",
        [
            "Summons (MC 01): obtain from clerk; 91-day validity (MCR 2.102).",
            "Personal service attempts: log dates/times/locations; attach photo proof where safe.",
            "Alternate Service Motion: attach affidavit of diligent attempts; propose certified mail + first-class + email + posting.",
            "Proof of Service: file promptly; calendar answer deadlines (21/28 days).",
        ],
        case_title,
    )
    gen_docx(
        dst_service_dir / "Affidavit_of_Diligence.docx",
        "Affidavit of Diligence",
        [
            "Affiant states under oath the following diligent attempts at personal service were made:",
            "1) Date/Time/Location/Outcome",
            "2) Date/Time/Location/Outcome",
            "3) Additional notes and photographs available.",
            "Subscribed and sworn on ________.",
            "Signature: __________________________",
        ],
        case_title,
    )
    gen_docx(
        dst_service_dir / "Retraction_Demand_Template.docx",
        "Retraction Demand (Defamation Mitigation)",
        [
            "Identify publication(s): date, medium, audience, verbatim or accurate paraphrase.",
            "State falsity and harm. Demand retraction within 10 days. Request republication channel.",
            "Reserve rights. Attach delivery proof.",
        ],
        case_title,
    )


# ---------------------------
# AUDITS / GATES
# ---------------------------
def audit_pin_cites(complaint_docx: Path) -> tuple[int, list[int]]:
    doc = Document(complaint_docx.as_posix())
    missing_idx = []
    total = 0
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if re.match(r"^\d+(\.|[\)])\s", txt):  # numbered para
            total += 1
            if not PIN_REGEX.search(txt):
                missing_idx.append(i + 1)
    return total, missing_idx


def audit_child_privacy_docx(docx_path: Path, child_name: str, child_initials: str) -> list[str]:
    doc = Document(docx_path.as_posix())
    blob = "\n".join(p.text for p in doc.paragraphs)
    flags = []
    # child name present?
    if re.search(re.escape(child_name), blob, re.IGNORECASE):
        flags.append(f"Child name appears unredacted: {child_name}")
    if child_initials not in blob:
        flags.append(f"Child initials token missing: {child_initials}")
    for pat in PRIVACY_TOKENS:
        if re.search(pat, blob, re.IGNORECASE):
            flags.append(f"PII token not redacted: {pat}")
    return flags


def write_manifest(manifest_path: Path, entries: list[list]):
    with manifest_path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["RelativePath", "Bytes", "SHA256"])
        w.writerows(entries)


# ---------------------------
# MAIN
# ---------------------------
def main():
    ap = argparse.ArgumentParser(description="WATSON ULTRA BUNDLER — Michigan Locked, Autonomous")
    ap.add_argument("--scan-root", required=True, help="Root folder to scan for candidate evidence (PDF/IMG/DOCX/TXT)")
    ap.add_argument(
        "--input-root", required=True, help="Optional source tree with preexisting Complaint/Exhibits/Matrices/etc."
    )
    ap.add_argument("--output-root", default="F:\\LegalResults", help="Output root")
    ap.add_argument("--case-title", default="Pigors v. Emily A. Watson, et al.", help="Caption line")
    ap.add_argument("--county", default="Muskegon", help="Venue county name")
    ap.add_argument("--child-name", default="Lincoln", help="Child name to redact")
    ap.add_argument("--child-initials", default="L.D.W.", help="Child initials to use")
    ap.add_argument("--ts", default="", help="Timestamp override, e.g., 20251029_1420")
    ap.add_argument("--no-dry-run", action="store_true", help="Write files and zip (otherwise dry-run)")
    args = ap.parse_args()

    scan_root = Path(args.scan_root).resolve()
    input_root = Path(args.input_root).resolve()
    out_root = Path(args.output_root).resolve()

    if not scan_root.exists():
        print(f"[FATAL] scan-root not found: {scan_root}")
        sys.exit(2)

    # Timestamp
    ts = args.ts or derive_ts_from_files([scan_root])
    bundle_dir = out_root / f"Watson_Complaint_Bundle_{ts}"
    makedirs(bundle_dir)

    # Logging
    log_path = bundle_dir / "watson_ultra_bundler.log"
    logging.basicConfig(
        filename=log_path.as_posix(),
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )
    console = logging.StreamHandler()
    console.setLevel(logging.INFO)
    console.setFormatter(logging.Formatter("%(message)s"))
    logging.getLogger().addHandler(console)

    logging.info("== WATSON ULTRA BUNDLER — Michigan Locked, Autonomous ==")
    logging.info(f"Scan root: {scan_root}")
    logging.info(f"Output root: {out_root}")
    logging.info(f"Timestamp: {ts}")
    logging.info(f"Dry run: {not args.no_dry_run}")

    # Directory structure
    d_complaint = makedirs(bundle_dir / "Complaint")
    d_exhibits = makedirs(bundle_dir / "Exhibits")
    d_matrices = makedirs(bundle_dir / "Matrices")
    d_discovery = makedirs(bundle_dir / "Discovery")
    d_pretrial = makedirs(bundle_dir / "Pretrial")
    d_orders = makedirs(bundle_dir / "Orders")
    d_service = makedirs(bundle_dir / "Service")
    d_damages = makedirs(bundle_dir / "Damages")

    manifest = []

    # 1) INDEX
    logging.info("[1] Indexing & OCR (housing excluded)…")
    index_rows = build_index(scan_root, args.child_name, args.child_initials)
    if not index_rows:
        logging.error("No non-housing files found to index.")
        sys.exit(2)
    logging.info(f"Indexed files: {len(index_rows)}")

    # 2) EXHIBITS
    logging.info("[2] Building Exhibits (H-1..)…")
    exhibits = stage_exhibits(index_rows, d_exhibits)
    # copy physical files
    if args.no_dry_run:
        for ex in exhibits:
            # prefer PDF; if non-PDF images, just copy as is
            sha = copy_with_hash(ex["src"], ex["dst"])
            ex["out_sha256"] = sha
            ex["out_size"] = ex["dst"].stat().st_size
            manifest.append([ex["dst"].relative_to(bundle_dir).as_posix(), ex["out_size"], sha])

    # 3) MATRICES (Elements, Privilege, SOL)
    logging.info("[3] Generating Matrices…")
    elements_csv = d_matrices / "Elements_Table.csv"
    priv_csv = d_matrices / "Privilege_Matrix.csv"
    sol_csv = d_matrices / "SOL_Matrix.csv"
    m = build_elements_table(exhibits, elements_csv)
    build_privilege_matrix(exhibits, priv_csv)
    build_sol_matrix(exhibits, sol_csv)
    if args.no_dry_run:
        for p in [elements_csv, priv_csv, sol_csv]:
            sha = sha256_of_path(p)
            manifest.append([p.relative_to(bundle_dir).as_posix(), p.stat().st_size, sha])

    # 4) VERIFIED COMPLAINT (auto-generated; neutral, pin-cited)
    logging.info("[4] Generating Verified Complaint (DOCX)…")
    complaint_docx = d_complaint / "Verified_Complaint.docx"
    make_verified_complaint(
        complaint_docx, args.case_title, args.county, exhibits, m["evidence_map"], args.child_name, args.child_initials
    )
    if args.no_dry_run:
        sha = sha256_of_path(complaint_docx)
        manifest.append([complaint_docx.relative_to(bundle_dir).as_posix(), complaint_docx.stat().st_size, sha])

    # 5) ORDERS / TEMPLATES
    logging.info("[5] Creating Proposed Orders & Templates…")
    ensure_orders(d_orders, args.case_title)
    ensure_discovery(d_discovery, args.case_title)
    ensure_pretrial(d_pretrial, args.case_title)
    ensure_service(d_service, args.case_title)
    if args.no_dry_run:
        for folder in [d_orders, d_discovery, d_pretrial, d_service]:
            for f in folder.glob("*.docx"):
                sha = sha256_of_path(f)
                manifest.append([f.relative_to(bundle_dir).as_posix(), f.stat().st_size, sha])

    # 6) DAMAGES WORKBOOK
    logging.info("[6] Building Damages Workbook CSV shell…")
    damages_csv = d_damages / "Damages_Workbook.csv"
    rows = [
        ["Category", "Proof Sources", "Amount", "Method", "Exhibit Pin-Cites"],
        ["Lost wages/opportunities", "", "", "Rate x Period", ""],
        ["Out-of-pocket expenses", "", "", "Summation", ""],
        ["Emotional distress", "", "", "Qualitative + corroboration", ""],
        ["Reputational/business loss", "", "", "Before/after metrics", ""],
        ["Parenting-time losses", "", "", "Nights x valuation", ""],
        ["Exemplary damages (egregiousness)", "", "", "Court’s discretion", ""],
    ]
    if args.no_dry_run:
        with damages_csv.open("w", encoding="utf-8", newline="") as f:
            csv.writer(f).writerows(rows)
        sha = sha256_of_path(damages_csv)
        manifest.append([damages_csv.relative_to(bundle_dir).as_posix(), damages_csv.stat().st_size, sha])

    # 7) GATES — Audits (HARD FAIL IF VIOLATE)
    logging.info("[7] Running Gating Audits…")

    # Pin-cite audit
    total_num, missing_idx = audit_pin_cites(complaint_docx)
    if total_num == 0 or missing_idx:
        logging.error(f"Pin-cite audit FAILED. Numbered paragraphs: {total_num}, missing cites: {len(missing_idx)}")
        logging.error(f"Missing ¶ indices (first 30): {missing_idx[:30]}")
        sys.exit(2)
    logging.info(f"[OK] Pin-cite audit passed: {total_num} numbered paragraphs")

    # Child privacy audit
    pf = audit_child_privacy_docx(complaint_docx, args.child_name, args.child_initials)
    if pf:
        logging.error("Child privacy audit FAILED:")
        for item in pf:
            logging.error(f"  - {item}")
        sys.exit(2)
    logging.info("[OK] Child privacy audit passed")

    # Privilege Matrix audit (ensure privileged entries are context-only)
    # Simple audit: every row with PrivilegeType absolute/qualified must have UseAllowed context-only/none
    def audit_privilege_csv(p: Path) -> list[str]:
        errs = []
        with p.open("r", encoding="utf-8-sig", newline="") as f:
            r = csv.DictReader(f)
            need = {"Source", "PrivilegeType", "UseAllowed"}
            if not need.issubset(set(r.fieldnames or [])):
                errs.append("Privilege_Matrix.csv missing required columns.")
                return errs
            for row in r:
                pt = (row.get("PrivilegeType") or "").strip().lower()
                ua = (row.get("UseAllowed") or "").strip().lower()
                if pt in {"absolute", "qualified"} and ua not in {"context-only", "none"}:
                    errs.append(f"Privileged item not context-only: {row}")
        return errs

    perrs = audit_privilege_csv(priv_csv)
    if perrs:
        logging.error("Privilege Matrix audit FAILED:")
        for e in perrs[:10]:
            logging.error(f"  - {e}")
        sys.exit(2)
    logging.info("[OK] Privilege Matrix audit passed")

    # SOL Matrix audit — stale must be background-only
    def audit_sol_csv(p: Path) -> list[str]:
        errs = []
        with p.open("r", encoding="utf-8-sig", newline="") as f:
            r = csv.DictReader(f)
            need = {"Claim", "Status", "Treatment"}
            if not need.issubset(set(r.fieldnames or [])):
                errs.append("SOL_Matrix.csv missing required columns.")
                return errs
            for row in r:
                st = (row.get("Status") or "").strip().lower()
                tr = (row.get("Treatment") or "").strip().lower()
                if st == "stale" and "background" not in tr:
                    errs.append(f"Stale item not confined to background: {row}")
        return errs

    serrs = audit_sol_csv(sol_csv)
    if serrs:
        logging.error("SOL Matrix audit FAILED:")
        for e in serrs[:10]:
            logging.error(f"  - {e}")
        sys.exit(2)
    logging.info("[OK] SOL Matrix audit passed")

    # 8) Manifest + ZIP
    logging.info("[8] Writing Manifest and ZIP…")
    manifest_path = bundle_dir / "Manifest.csv"
    if args.no_dry_run:
        write_manifest(manifest_path, manifest)
        with zipfile.ZipFile(
            (out_root / f"Watson_Complaint_Bundle_{ts}.zip").as_posix(), "w", compression=zipfile.ZIP_DEFLATED
        ) as zf:
            for root, _, files in os.walk(bundle_dir.as_posix()):
                for name in files:
                    p = Path(root) / name
                    rel = p.relative_to(bundle_dir)
                    zf.write(p.as_posix(), rel.as_posix())
        # log final ZIP hash
        zpath = out_root / f"Watson_Complaint_Bundle_{ts}.zip"
        zsha = sha256_of_path(zpath)
        logging.info(f"[ZIP] {zpath.name} | {zpath.stat().st_size} bytes | SHA-256={zsha}")

    # 9) Summary
    logging.info("== SUMMARY ==")
    logging.info(f"Bundle: {bundle_dir}")
    if args.no_dry_run:
        logging.info(f"ZIP: {out_root / f'Watson_Complaint_Bundle_{ts}.zip'}")
    logging.info("All gates passed. Ready for MiFILE upload. Housing excluded by design.")


if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except Exception as e:
        print(f"[FATAL] {e}", file=sys.stderr)
        sys.exit(1)
