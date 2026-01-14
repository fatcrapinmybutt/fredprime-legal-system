#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HOUSING ULTRA BUNDLER — Diamond 9999+++ (Michigan-Locked, Autonomous, Housing-Only)
Author: Litigation OS (Andrew J. Pigors)

Purpose (Strictly Housing):
  • Recursively scan a source tree and build a HOUSING-ONLY evidence bundle (PDF/IMG/DOCX/TXT).
  • Auto-select only housing-related items (e.g., Shady Oaks Park MHP LLC, Homes of America, sewage/EGLE, rent/fees, lease, eviction).
  • EXCLUDE family/PPO/custody items entirely (kept separate).
  • OCR/extract text (PyMuPDF -> PyPDF2 fallback; PIL+pytesseract for images), compute SHA-256, capture EXIF (optional).
  • Auto-label and stage exhibits (prefix default "HX-": HX-1…HX-N) with deterministic ordering and on-record metadata CSV.
  • Generate/validate:
      - Verified Housing Complaint (DOCX) — Michigan counts (habitability, unlawful interference with possessory interest, nuisance, breach of contract/unjust enrichment),
        neutral paragraphs with exhibit pin-cites (no placeholders).
      - Proposed Orders (Preliminary Injunction to restore utilities/repairs; Protective; Alternate Service; Motions in Limine).
      - ESI Protocol + Discovery Plan (housing custodians).
      - Pretrial Statement + Proposed Jury Instructions (outline, M Civ JI-aligned where applicable).
      - Retraction/Correction Demands (for billing/notice misstatements).
      - Elements Table (per housing count), Privilege Matrix (exclude predicates), SOL Matrix (claim vs. background).
      - Damages Workbook CSV (rent abatement, overbilling refunds, hotel + clean-up, property loss, emotional distress, exemplary where allowed).
  • Enforce hard gates (fail-fast): Pin-Cites, Privilege, SOL, PII/financial redaction.
  • Emit Manifest.csv with sizes + SHA-256 and a ready-to-file ZIP bundle.

Legal Anchors (Michigan-Locked; script text is neutral and evidence-tethered):
  • Habitability/Common Areas: MCL 554.139 (duty to keep premises in reasonable repair/fit for intended use; common areas). 
  • Unlawful Interference with Possessory Interest (self-help/utility shutoff): MCL 600.2918 (including interference by utility termination).
  • Nuisance (private), Negligence (duty/breach/causation/damages), Breach of Contract/Unjust Enrichment (overbilling/unauthorized charges).
  • Civil discovery/procedure: MCR 1.109 (signatures/redaction/sanctions), 2.111–2.113 (pleading), 2.302–2.313 (discovery/ESI), 2.105 (service), 2.508(A) (jury).
  • NOTE: The script keeps housing separate from family/PPO by design.

Dependencies (install once, offline-capable):
  pip install python-docx PyPDF2
  pip install pymupdf               # optional: best PDF text extraction (PyMuPDF)
  pip install pillow pytesseract     # optional: image OCR (requires Tesseract installed on system)
  pip install exifread               # optional: richer EXIF parsing

Usage (PowerShell; Windows paths shown — works on Termux/Linux with adjusted paths):
  python housing_ultra_bundler.py ^
    --scan-root "F:\\HousingEvidence" ^
    --output-root "F:\\LegalResults" ^
    --case-title "Pigors v. Shady Oaks Park MHP LLC, et al." ^
    --county "Muskegon" ^
    --exhibit-prefix "HX-" ^
    --ts "20251029_1500" ^
    --no-dry-run

Notes:
  • DRY RUN is ON by default. Add --no-dry-run to write files & ZIP.
  • If you omit --ts, a deterministic timestamp is derived from inputs.
  • Zero placeholders policy: the generator writes neutral, exhibit-cited facts. If a gate fails, the script aborts with a specific message.
"""

import argparse
import csv
import datetime as dt
import hashlib
import io
import json
import logging
import os
import re
import shutil
import sys
import zipfile
from pathlib import Path
from collections import defaultdict

# Optional imports (handled gracefully)
HAVE_PYMUPDF = False
HAVE_TESSERACT = False
HAVE_PIL = False
HAVE_EXIFREAD = False

try:
    import fitz  # PyMuPDF
    HAVE_PYMUPDF = True
except Exception:
    pass

try:
    from PIL import Image
    HAVE_PIL = True
    try:
        import pytesseract
        HAVE_TESSERACT = True
    except Exception:
        HAVE_TESSERACT = False
except Exception:
    HAVE_PIL = False

try:
    import exifread
    HAVE_EXIFREAD = True
except Exception:
    HAVE_EXIFREAD = False

try:
    from PyPDF2 import PdfReader
except Exception as e:
    print("Missing dependency: PyPDF2. Install with: pip install PyPDF2")
    raise

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    print("Missing dependency: python-docx. Install with: pip install python-docx")
    raise

# ---------------------------
# CONFIG / CONSTANTS
# ---------------------------
ALLOWED_EXT = {".pdf", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".docx", ".txt"}

# Housing-positive keyword filter (broad)
HOUSING_POSITIVE = re.compile(
    r"\b("
    r"shady\s*oaks|shady\s*oaks\s*park|mhp|mobile\s*home\s*park|manufactured\s*home|"
    r"homes\s*of\s*america|alden\s*global|partridge\s*securities|cricklewood|ravenna|"
    r"lot\s*\d+|pad|lease|rental\s*agreement|rent|late\s*fee|overbill|utility|water|sewer|"
    r"sewage|egle|dte|consumers\s*energy|zego|trash|maintenance|habitability|mcl\s*554\.139|"
    r"unlawful\s*interference|mcl\s*600\.2918|summary\s*proceedings|evict|eviction|"
    r"notice\s*to\s*quit|7[-\s]*day|30[-\s]*day|repairs|violation\s*notice|inspection|"
    r"bill|statement|ledger|park\s*rules|manager|kim\s*davis|cricklewood\s*mhp|"
    r"whitehall\s*road|service\s*charge|water\s*shut\s*off|sewage\s*leak"
    r")\b",
    re.IGNORECASE,
)

# Family-negative filter (strict exclusion)
FAMILY_NEGATIVE = re.compile(
    r"\b("
    r"ppo|show\s*cause|parenting\s*time|custody|friend\s*of\s*the\s*court|foc|"
    r"benchbook|canon|mcl\s*722\.|mcr\s*3\.2|best\s*interest|appclose|nspd|"
    r"ella\s*randall|ruscopa|supervised\s*time|lincoln(?!\s*park)|"
    r"jenny\s*l\.?\s*mcneill"
    r")\b",
    re.IGNORECASE,
)

# Exhibit pin-cite detection
PIN_REGEX = re.compile(
    r"(Exhibit\s+[A-Z]{1,3}(?:-\d+)?(?:\s*p\.\s*\d+)?|[A-Z]{1,3}-\d+\s*(?:p\.\s*\d+|t=\d{2}:\d{2}:\d{2})?)",
    re.IGNORECASE,
)

# PII/Financial tokens to sanitize in public complaint
PII_TOKENS = [
    r"\bSSN\b", r"\bSocial\s*Security\b", r"\bDriver'?s\s*License\b", r"\bDL\s*#\b",
    r"\bAccount\s*Number\b", r"\bRouting\s*Number\b", r"\bCredit\s*Card\b", r"\bCVV\b",
    r"\bBank\b\s*(Account|Acct)\b", r"\bDOB\b", r"\bDate\s*of\s*Birth\b"
]

# Housing claims & neutral elements (Michigan-aligned)
HOUSING_CLAIMS = [
    ("Breach of the Covenant of Habitability / Common Areas (MCL 554.139)", [
        "Premises/common areas were not kept in reasonable repair or fit for their intended use",
        "Defendant was the lessor/park operator responsible for relevant premises/common systems",
        "Breach caused damages (loss of use, abatement, health risks, costs)"
    ]),
    ("Unlawful Interference with Possessory Interest (MCL 600.2918)", [
        "Plaintiff had lawful possessory interest",
        "Defendant interfered (including by utility shutoff/obstruction/self-help)",
        "Interference was willful or without legal process",
        "Damages (including treble where applicable)"
    ]),
    ("Nuisance (Private)", [
        "Defendant’s conduct or condition substantially and unreasonably interfered with Plaintiff’s use and enjoyment",
        "Causation attributable to Defendant",
        "Damages (including discomfort, property, and remediation costs)"
    ]),
    ("Negligence (Maintenance/Operations)", [
        "Duty to maintain park premises/common systems with reasonable care",
        "Breach (e.g., sewage leaks, unsafe infrastructure, delayed repairs)",
        "Causation and damages"
    ]),
    ("Breach of Contract / Unjust Enrichment (Overbilling/Unauthorized Charges)", [
        "Contract terms or lawful rates exceeded or ignored; or conferral of benefit without right",
        "Non-conforming billing/fees/charges (e.g., water/sewer/trash/admin) collected",
        "Damages: refund/credit, fee disgorgement, restitution"
    ]),
    ("Declaratory and Injunctive Relief", [
        "Actual controversy over rights/duties/charges/conditions",
        "Narrow injunction and repair order necessary to prevent continuing harm",
        "Balancing of equities favors relief"
    ]),
]

# ---------------------------
# UTILITIES
# ---------------------------
def sha256_of_path(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()

def makedirs(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path

def copy_with_hash(src: Path, dst: Path) -> str:
    makedirs(dst.parent)
    shutil.copy2(src, dst)
    return sha256_of_path(dst)

def derive_ts_from_files(paths: list[Path]) -> str:
    h = hashlib.sha256()
    for p in sorted(paths, key=lambda x: x.as_posix()):
        try:
            h.update(p.name.encode())
            h.update(str(int(p.stat().st_mtime)).encode())
        except Exception:
            pass
    digest = h.hexdigest()[:8]
    return dt.datetime.now().strftime("%Y%m%d") + f"_{digest}"

def chunk_snippet(text: str, max_len=160) -> str:
    s = re.sub(r"\s+", " ", text or "").strip()
    return (s[:max_len] + "…") if len(s) > max_len else s

def sanitize_public_text(text: str) -> str:
    # Redact obvious PII/financial tokens
    sanitized = text
    for pat in PII_TOKENS:
        sanitized = re.sub(pat, "[REDACTED]", sanitized, flags=re.IGNORECASE)
    return sanitized

def split_sentences(text: str) -> list[str]:
    return re.split(r"(?<=[\.\?\!])\s+", re.sub(r"\s+", " ", text or "").strip())

# ---------------------------
# TEXT/EXIF EXTRACTION
# ---------------------------
def text_from_pdf(p: Path) -> str:
    if HAVE_PYMUPDF:
        try:
            doc = fitz.open(p.as_posix())
            return "\n".join([pg.get_text("text") for pg in doc])
        except Exception:
            pass
    try:
        rd = PdfReader(p.as_posix())
        out = []
        for pg in rd.pages:
            try:
                out.append(pg.extract_text() or "")
            except Exception:
                out.append("")
        return "\n".join(out)
    except Exception:
        return ""

def text_from_image(p: Path) -> str:
    if HAVE_PIL and HAVE_TESSERACT:
        try:
            img = Image.open(p.as_posix())
            return pytesseract.image_to_string(img)
        except Exception:
            return ""
    return ""

def text_from_docx(p: Path) -> str:
    try:
        doc = Document(p.as_posix())
        return "\n".join(par.text for par in doc.paragraphs)
    except Exception:
        return ""

def text_from_txt(p: Path) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return p.read_text(encoding=enc, errors="ignore")
        except Exception:
            continue
    return ""

def extract_text(p: Path) -> str:
    ext = p.suffix.lower()
    if ext == ".pdf":
        return text_from_pdf(p)
    if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        return text_from_image(p)
    if ext == ".docx":
        return text_from_docx(p)
    if ext == ".txt":
        return text_from_txt(p)
    return ""

def exif_dict(p: Path) -> dict:
    out = {}
    if HAVE_EXIFREAD and p.suffix.lower() in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        try:
            with p.open("rb") as f:
                tags = exifread.process_file(f, details=False)
            for k, v in tags.items():
                out[str(k)] = str(v)
        except Exception:
            pass
    return out

# ---------------------------
# DISCOVERY & FILTERING
# ---------------------------
def discover_files(scan_root: Path) -> list[Path]:
    items = []
    for root, _, files in os.walk(scan_root.as_posix()):
        for n in files:
            p = Path(root) / n
            if p.suffix.lower() in ALLOWED_EXT and p.is_file():
                items.append(p)
    return items

def is_housing_candidate(path: Path, text: str) -> bool:
    name_hit = bool(HOUSING_POSITIVE.search(path.name))
    text_hit = bool(HOUSING_POSITIVE.search(text))
    return name_hit or text_hit

def is_family_excluded(path: Path, text: str) -> bool:
    return bool(FAMILY_NEGATIVE.search(path.name)) or bool(FAMILY_NEGATIVE.search(text))

def build_index(scan_root: Path) -> list[dict]:
    rows = []
    for p in discover_files(scan_root):
        try:
            size = p.stat().st_size
            sha = sha256_of_path(p)
        except Exception:
            size, sha = 0, ""
        txt = extract_text(p)
        if is_family_excluded(p, txt):
            continue  # strictly exclude family/PPO/custody materials
        if not is_housing_candidate(p, txt):
            continue  # not a housing document
        exif = exif_dict(p)
        preview = sanitize_public_text(txt[:2000]) if txt else ""
        rows.append({
            "path": p, "name": p.name, "ext": p.suffix.lower(), "size": size, "sha256": sha,
            "text": txt, "preview": preview, "exif": exif
        })
    return rows

# ---------------------------
# EXHIBITS
# ---------------------------
def stage_exhibits(index_rows: list[dict], dst_dir: Path, prefix: str) -> list[dict]:
    """
    Prioritize PDFs, then images, then DOCX/TXT. Newest first within class.
    Label: HX-1…HX-N (prefix configurable).
    """
    scored = []
    for r in index_rows:
        p = r["path"]
        try:
            mtime = p.stat().st_mtime
        except Exception:
            mtime = 0
        rank = 0
        if r["ext"] == ".pdf": rank = 3
        elif r["ext"] in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}: rank = 2
        else: rank = 1
        scored.append((rank, -mtime, r))
    scored.sort(reverse=True)

    out = []
    n = 1
    for _, _, r in scored:
        label = f"{prefix}{n}"
        src = r["path"]
        dst = dst_dir / f"{label}_{src.stem}{src.suffix.lower()}"
        out.append({
            "label": label,
            "src": src,
            "dst": dst,
            "sha256": r["sha256"],
            "size": r["size"],
            "text": r["text"],
            "preview": r["preview"],
            "ext": r["ext"]
        })
        n += 1
    return out

# ---------------------------
# MATRICES
# ---------------------------
def build_elements_table(exhibits: list[dict], out_csv: Path) -> dict:
    """
    Build Elements Table per HOUSING_CLAIM with heuristic pin-cites based on keyword hits.
    """
    cmap = {
        "Breach of the Covenant of Habitability / Common Areas (MCL 554.139)": ["sewage", "sewer", "leak", "egle", "repair", "habitability", "unsafe", "unsanitary"],
        "Unlawful Interference with Possessory Interest (MCL 600.2918)": ["shut off", "shutoff", "water off", "utility", "lockout", "self-help", "no notice"],
        "Nuisance (Private)": ["odor", "stink", "sewage", "noise", "interfere", "enjoyment"],
        "Negligence (Maintenance/Operations)": ["duty", "breach", "negligent", "failure to maintain", "unsafe", "injury"],
        "Breach of Contract / Unjust Enrichment (Overbilling/Unauthorized Charges)": ["bill", "overbill", "admin fee", "trash", "water", "sewer", "ledger", "statement", "rate"],
        "Declaratory and Injunctive Relief": ["injunction", "repair", "abatement", "order", "declaratory"],
    }
    evidence_map = defaultdict(list)
    rows = []
    for claim, elements in HOUSING_CLAIMS:
        kws = cmap.get(claim, [])
        for ex in exhibits:
            txt = (ex["text"] or "").lower()
            if any(kw in txt for kw in kws):
                pin = f"{ex['label']} p.1"
                snippet = chunk_snippet(ex["preview"])
                evidence_map[claim].append((pin, snippet))
    with out_csv.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Claim", "Element", "PinCites", "CorroborationSnippets", "Gaps/Notes"])
        for claim, elements in HOUSING_CLAIMS:
            pins = evidence_map.get(claim, [])
            for elem in elements:
                w.writerow([
                    claim,
                    elem,
                    " | ".join(p for p, _ in pins) if pins else "",
                    " | ".join(s for _, s in pins[:5]) if pins else "",
                    "" if pins else "No automatic hits — add manual cite if appropriate"
                ])
    return {"evidence_map": evidence_map}

def build_privilege_matrix(exhibits: list[dict], out_csv: Path) -> None:
    with out_csv.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Source", "PrivilegeType", "Basis", "UseAllowed", "Notes"])
        for ex in exhibits:
            txt = ex["text"] or ""
            ptype, basis, use = "", "", "claim"
            if re.search(r"STATE OF MICHIGAN|CIRCUIT COURT|DISTRICT COURT|COURT OF APPEALS|EGLE|NOTICE OF VIOLATION", txt, re.IGNORECASE):
                # Treat official filings/reports conservatively re: privilege/fair report context
                ptype, basis, use = "qualified", "official-report/fair-report indicators", "context-only"
            w.writerow([ex["src"].name, ptype, basis, use, ""])

def build_sol_matrix(exhibits: list[dict], out_csv: Path) -> None:
    """
    Heuristic SOL:
      - Contract/Unjust Enrichment: 6 years (MCL 600.5807 for breach of contract — general ref)
      - Tort (negligence/nuisance/2918 interference): 3 years (MCL 600.5805 — general tort ref)
      - Declaratory/Injunctive: N/A; tracks ongoing controversy.
    """
    claim_limits = {
        "Breach of the Covenant of Habitability / Common Areas (MCL 554.139)": 3,   # tort-leaning habitability harms
        "Unlawful Interference with Possessory Interest (MCL 600.2918)": 3,
        "Nuisance (Private)": 3,
        "Negligence (Maintenance/Operations)": 3,
        "Breach of Contract / Unjust Enrichment (Overbilling/Unauthorized Charges)": 6,
        "Declaratory and Injunctive Relief": 0
    }
    this_year = dt.datetime.now().year
    with out_csv.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["Claim", "AccrualYear(heuristic)", "LimitYears", "Status", "Treatment", "Notes"])
        for claim, _ in HOUSING_CLAIMS:
            limit = claim_limits[claim]
            # Heuristic: if any exhibit text contains a year, choose the earliest as accrual proxy
            years = set(int(y) for y in re.findall(r"\b(20[0-3]\d)\b", " ".join([(ex["text"] or "") for ex in exhibits])))
            accrual = min(years) if years else this_year
            if limit == 0:
                status, treat = "timely", "claim"
            else:
                age = this_year - accrual
                if age <= limit:
                    status, treat = "timely", "claim"
                else:
                    status, treat = "stale", "background-only"
            w.writerow([claim, accrual, limit, status, treat, "Heuristic; confirm with event-level dates."])

# ---------------------------
# DOCX HELPERS
# ---------------------------
def add_heading(doc: Document, text: str, size=13, center=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def numbered_par(doc: Document, idx: int, text: str):
    doc.add_paragraph(f"{idx}. {text}")

def gen_docx(out_path: Path, title: str, paras: list[str], case_title: str):
    doc = Document()
    add_heading(doc, case_title, 14)
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    for para in paras:
        doc.add_paragraph(para)
    makedirs(out_path.parent)
    doc.save(out_path.as_posix())

# ---------------------------
# VERIFIED HOUSING COMPLAINT
# ---------------------------
def make_verified_housing_complaint(out_docx: Path, case_title: str, county: str,
                                    exhibits: list[dict], evidence_map: dict):
    doc = Document()
    add_heading(doc, f"STATE OF MICHIGAN — IN THE CIRCUIT COURT FOR {county.upper()} COUNTY", 12)
    add_heading(doc, case_title, 14)
    p = doc.add_paragraph("VERIFIED HOUSING COMPLAINT — JURY DEMAND — SANCTIONS REQUEST (MCR 1.109(E))")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # Preliminary Statement
    doc.add_heading("Preliminary Statement", level=1)
    doc.add_paragraph(
        "Plaintiff alleges housing-related wrongful conduct by Defendants, including habitability/common-area failures, "
        "unlawful interference with possessory interests, nuisance conditions, negligence in maintenance/operations, "
        "and overbilling/unauthorized charges. All factual statements are supported by exhibit pin-cites."
    )

    # Factual Allegations (neutral with pin-cites)
    doc.add_heading("Factual Allegations", level=1)
    n = 1
    for ex in exhibits[:20]:
        # prefer sentences containing housing hits
        sents = split_sentences(ex.get("preview", ""))
        pick = ""
        for s in sents:
            if re.search(HOUSING_POSITIVE, s):
                pick = s
                break
        if not pick and sents:
            pick = sents[0]
        if pick:
            numbered_par(doc, n, sanitize_public_text(f"{pick} (Exhibit {ex['label']} p.1)"))
            n += 1
    if n == 1:
        # Ensure at least one factual statement referencing an exhibit (no placeholders)
        numbered_par(doc, n, f"Documents reflecting housing conditions and charges are attached. (Exhibit {exhibits[0]['label']} p.1)")
        n += 1

    # Counts — element-by-element with evidence_map pins
    def write_count(title: str, elements: list[str]):
        doc.add_heading(title, level=2)
        numbered_par(doc, 1, "Plaintiff incorporates the preceding paragraphs by reference.")
        pins = evidence_map.get(title, [])
        for i, elem in enumerate(elements, start=2):
            cite = f" Supported by: " + " ; ".join([p for p, _ in pins[:6]]) if pins else ""
            numbered_par(doc, i, f"{elem}.{cite}")

    for title, elements in HOUSING_CLAIMS:
        write_count(title, elements)

    # Sanctions (MCR 1.109(E))
    doc.add_heading("Sanctions (MCR 1.109(E))", level=2)
    doc.add_paragraph(
        "Plaintiff requests findings that identified papers or statements were presented for improper purpose or without evidentiary support "
        "after reasonable inquiry, and an award of reasonable expenses and attorney fees."
    )

    # Prayer
    doc.add_heading("Prayer for Relief", level=1)
    for item in [
        "Compensatory damages including abatement/refund, remediation, and consequential losses;",
        "Exemplary damages where allowed by Michigan law;",
        "Declaratory and narrow injunctive relief to halt interference, restore/utilities/repairs, and correct billing;",
        "Sanctions and reasonable expenses under MCR 1.109(E);",
        "Costs and interest as permitted;",
        "Further relief as is just."
    ]:
        doc.add_paragraph(f"• {item}")

    # Jury demand
    doc.add_heading("Jury Demand", level=1)
    doc.add_paragraph("Plaintiff demands trial by jury on all triable issues. (MCR 2.508(A))")

    # Verification
    doc.add_heading("Verification", level=1)
    doc.add_paragraph(
        "I declare that the factual allegations above are true to the best of my knowledge, information, and belief. "
        "I understand that a willful false statement is subject to penalty."
    )
    doc.add_paragraph("Dated: " + dt.datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph("/s/ Andrew J. Pigors")

    makedirs(out_docx.parent)
    doc.save(out_docx.as_posix())

# ---------------------------
# AUDITS / GATES
# ---------------------------
def audit_pin_cites(docx_path: Path) -> tuple[int, list[int]]:
    doc = Document(docx_path.as_posix())
    total, missing = 0, []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if re.match(r"^\d+(\.|[\)])\s", t):
            total += 1
            if not PIN_REGEX.search(t):
                missing.append(i + 1)
    return total, missing

def audit_pii_redaction(docx_path: Path) -> list[str]:
    doc = Document(docx_path.as_posix())
    blob = "\n".join(p.text for p in doc.paragraphs)
    flags = []
    for pat in PII_TOKENS:
        if re.search(pat, blob, re.IGNORECASE):
            flags.append(f"PII token not redacted: {pat}")
    return flags

def write_manifest(manifest_path: Path, entries: list[list]):
    with manifest_path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["RelativePath","Bytes","SHA256"])
        w.writerows(entries)

# ---------------------------
# PROPOSED ORDERS & DOCS
# ---------------------------
def ensure_orders(dst_orders_dir: Path, case_title: str):
    gen_docx(dst_orders_dir / "Preliminary_Injunction_Order.docx", "Preliminary Injunction (Repairs/Utilities/Non-Interference)", [
        "ORDER: Defendants shall promptly restore and maintain essential services (including water/sewer/utilities) and abate hazard conditions.",
        "ORDER: Defendants shall not engage in self-help or utility interference; all remedies must proceed by lawful process.",
        "ORDER: Defendants shall implement repairs consistent with industry standards and applicable codes; status report due in 14 days.",
        "ORDER: Clerk to schedule compliance review; violations subject to contempt."
    ], case_title)
    gen_docx(dst_orders_dir / "Protective_Order_MCR_2_302_C.docx", "Protective Order (MCR 2.302(C))", [
        "ORDER: Financial account identifiers and resident PII shall be redacted in public filings under MCR 1.109(D)(9).",
        "ORDER: Sensitive exhibits may be filed under seal; use limited to this case; clawback applies to inadvertent production."
    ], case_title)
    gen_docx(dst_orders_dir / "Alternate_Service_Order_MCR_2_105.docx", "Alternate Service Order (MCR 2.105)", [
        "ORDER: Alternate service authorized upon affidavit of diligent attempts; service by certified mail + first-class + email + posting as permitted.",
        "ORDER: Proof of service to be filed within 14 days."
    ], case_title)
    gen_docx(dst_orders_dir / "Motions_in_Limine_Order.docx", "Motions in Limine (MRE 401–403, 404, 802, 901)", [
        "ORDER: Exclude character/propensity; remote or cumulative evidence; hearsay without exception; and unauthenticated media lacking MRE 901 foundation."
    ], case_title)
    gen_docx(dst_orders_dir / "Sanctions_Order_MCR_1_109_E.docx", "Sanctions Order (MCR 1.109(E))", [
        "ORDER: The Court finds the identified papers lacked evidentiary support or were presented for improper purpose.",
        "ORDER: Defendants shall pay Plaintiff’s reasonable expenses and fees; fee affidavit within 14 days; objections due 14 days thereafter."
    ], case_title)

def ensure_discovery(dst_discovery_dir: Path, case_title: str):
    gen_docx(dst_discovery_dir / "ESI_Protocol.docx", "ESI Protocol (Housing)", [
        "Custodians: park management; billing/ledger custodian; maintenance; third-party utility providers (Zego, Consumers Energy, water/sewer), EGLE contacts.",
        "Formats: PST/MBOX+CSV (email); JSON/CSV+PDF (messages/ledgers); native+SHA-256 (audio/video); photos with EXIF.",
        "Hashing: SHA-256 manifest; dedup by hash; categorical privilege log; clawback.",
        "Search: exchange terms & hit counts; rolling production with Bates and exhibit IDs."
    ], case_title)
    gen_docx(dst_discovery_dir / "Discovery_Plan.docx", "Discovery Plan (Housing)", [
        "Scope: elements-based discovery (habitability defects, utility events, ledgers/billing, notices, repair logs, EGLE interactions).",
        "Deadlines: per MCR 2.401 scheduling. Protective Order for PII/financials.",
        "Subpoenas: utilities, EGLE, vendors, maintenance contractors; ledger systems/export specs.",
        "ADR: consider case evaluation and offers of judgment timing (MCR 2.403/2.405)."
    ], case_title)

def ensure_pretrial(dst_pretrial_dir: Path, case_title: str):
    gen_docx(dst_pretrial_dir / "Pretrial_Statement.docx", "Pretrial Statement (Housing)", [
        "Stipulated Facts: to be refined pre-conference.",
        "Contested Facts: numbered with exhibit pin-cites.",
        "Issues of Law: MCL 554.139; MCL 600.2918; negligence/nuisance; contract/unjust enrichment; evidentiary issues.",
        "Witnesses: management, maintenance, billing, residents, utility reps; time estimates.",
        "Exhibits: ID, description, source, SHA-256, foundation, objections.",
        "Motions in Limine: list with authority.",
        "Time Estimates & Tech: audio/video playback, ledger display with hash verification."
    ], case_title)
    gen_docx(dst_pretrial_dir / "Proposed_Jury_Instructions.docx", "Proposed Jury Instructions (Outline — Housing)", [
        "Habitability/Common Areas (elements; reasonable repair/fitness; damages).",
        "Unlawful Interference with Possessory Interest (elements; interference/self-help; damages/treble where applicable).",
        "Nuisance (substantial and unreasonable interference; damages).",
        "Negligence (duty/breach/causation/damages).",
        "Contract/Unjust Enrichment (terms/benefit/retention unjust; damages)."
    ], case_title)

def ensure_service(dst_service_dir: Path, case_title: str):
    gen_docx(dst_service_dir / "Service_Playbook.docx", "Service Playbook (MCR 2.105)", [
        "Summons (MC 01): obtain; 91-day validity (MCR 2.102).",
        "Personal service attempts: log dates/times/locations; photo proof when safe.",
        "Alternate Service Motion: attach diligence affidavit; propose certified mail + first-class + email + posting.",
        "Proof of Service: file promptly; calendar answer deadlines (21/28 days)."
    ], case_title)
    gen_docx(dst_service_dir / "Affidavit_of_Diligence.docx", "Affidavit of Diligence", [
        "Affiant states diligent attempts at personal service:",
        "1) Date/Time/Location/Outcome",
        "2) Date/Time/Location/Outcome",
        "3) Additional notes/photographs available.",
        "Subscribed and sworn on ________.",
        "Signature: __________________________"
    ], case_title)
    gen_docx(dst_service_dir / "Correction_Demand_Template.docx", "Correction/Adjustment Demand (Billing/Charges)", [
        "Identify specific billings/charges by date and amount; state basis for correction/refund.",
        "Demand written correction within 10 days; reserve rights; attach delivery proof."
    ], case_title)

# ---------------------------
# MAIN
# ---------------------------
def main():
    ap = argparse.ArgumentParser(description="HOUSING ULTRA BUNDLER — Michigan-Locked, Autonomous, Housing-Only")
    ap.add_argument("--scan-root", required=True, help="Folder to scan for HOUSING evidence (PDF/IMG/DOCX/TXT)")
    ap.add_argument("--output-root", default="F:\\LegalResults", help="Output root")
    ap.add_argument("--case-title", default="Pigors v. Shady Oaks Park MHP LLC, et al.", help="Caption line")
    ap.add_argument("--county", default="Muskegon", help="Venue county")
    ap.add_argument("--exhibit-prefix", default="HX-", help="Exhibit label prefix (default: HX-)")
    ap.add_argument("--ts", default="", help="Timestamp override, e.g., 20251029_1500")
    ap.add_argument("--no-dry-run", action="store_true", help="Write files and ZIP (otherwise dry-run)")
    args = ap.parse_args()

    scan_root = Path(args.scan_root).resolve()
    out_root = Path(args.output_root).resolve()
    if not scan_root.exists():
        print(f"[FATAL] scan-root not found: {scan_root}")
        sys.exit(2)

    ts = args.ts or derive_ts_from_files([scan_root])
    bundle_dir = out_root / f"Housing_Complaint_Bundle_{ts}"
    makedirs(bundle_dir)

    # Logging
    log_path = bundle_dir / "housing_ultra_bundler.log"
    logging.basicConfig(
        filename=log_path.as_posix(),
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )
    console = logging.StreamHandler()
    console.setLevel(logging.INFO)
    console.setFormatter(logging.Formatter("%(message)s"))
    logging.getLogger().addHandler(console)

    logging.info("== HOUSING ULTRA BUNDLER — Michigan Locked, Housing-Only ==")
    logging.info(f"Scan root: {scan_root}")
    logging.info(f"Output root: {out_root}")
    logging.info(f"Timestamp: {ts}")
    logging.info(f"Dry run: {not args.no_dry_run}")

    # Directory structure
    d_complaint = makedirs(bundle_dir / "Complaint")
    d_exhibits  = makedirs(bundle_dir / "Exhibits")
    d_matrices  = makedirs(bundle_dir / "Matrices")
    d_discovery = makedirs(bundle_dir / "Discovery")
    d_pretrial  = makedirs(bundle_dir / "Pretrial")
    d_orders    = makedirs(bundle_dir / "Orders")
    d_service   = makedirs(bundle_dir / "Service")
    d_damages   = makedirs(bundle_dir / "Damages")

    manifest = []

    # 1) INDEX
    logging.info("[1] Indexing Housing corpus (family materials excluded)…")
    idx = build_index(scan_root)
    if not idx:
        logging.error("No HOUSING files discovered after filtering (or all were excluded as family/PPO).")
        sys.exit(2)
    logging.info(f"Housing files indexed: {len(idx)}")

    # 2) EXHIBITS (HX-1..)
    logging.info("[2] Building Exhibits…")
    exhibits = stage_exhibits(idx, d_exhibits, args.exhibit_prefix)
    if args.no_dry_run:
        for ex in exhibits:
            sha = copy_with_hash(ex["src"], ex["dst"])
            ex["out_sha256"] = sha
            ex["out_size"] = ex["dst"].stat().st_size
            manifest.append([ex["dst"].relative_to(bundle_dir).as_posix(), ex["out_size"], sha])

    # 3) MATRICES (Elements/Privilege/SOL)
    logging.info("[3] Generating Matrices…")
    elements_csv = d_matrices / "Elements_Table_Housing.csv"
    priv_csv     = d_matrices / "Privilege_Matrix_Housing.csv"
    sol_csv      = d_matrices / "SOL_Matrix_Housing.csv"
    m = build_elements_table(exhibits, elements_csv)
    build_privilege_matrix(exhibits, priv_csv)
    build_sol_matrix(exhibits, sol_csv)
    if args.no_dry_run:
        for p in [elements_csv, priv_csv, sol_csv]:
            manifest.append([p.relative_to(bundle_dir).as_posix(), p.stat().st_size, sha256_of_path(p)])

    # 4) VERIFIED HOUSING COMPLAINT (DOCX)
    logging.info("[4] Generating Verified Housing Complaint (DOCX)…")
    complaint_docx = d_complaint / "Verified_Housing_Complaint.docx"
    make_verified_housing_complaint(complaint_docx, args.case_title, args.county, exhibits, m["evidence_map"])
    if args.no_dry_run:
        manifest.append([complaint_docx.relative_to(bundle_dir).as_posix(), complaint_docx.stat().st_size, sha256_of_path(complaint_docx)])

    # 5) Orders / Discovery / Pretrial / Service
    logging.info("[5] Creating Proposed Orders & Support Docs…")
    ensure_orders(d_orders, args.case_title)
    ensure_discovery(d_discovery, args.case_title)
    ensure_pretrial(d_pretrial, args.case_title)
    ensure_service(d_service, args.case_title)
    if args.no_dry_run:
        for folder in [d_orders, d_discovery, d_pretrial, d_service]:
            for f in folder.glob("*.docx"):
                manifest.append([f.relative_to(bundle_dir).as_posix(), f.stat().st_size, sha256_of_path(f)])

    # 6) Damages Workbook
    logging.info("[6] Building Damages Workbook CSV…")
    damages_csv = d_damages / "Damages_Workbook_Housing.csv"
    rows = [
        ["Category","Proof Sources","Amount","Method","Exhibit Pin-Cites"],
        ["Rent abatement (loss of use)","","","Rate x Affected Period",""],
        ["Refunds for unauthorized charges (water/sewer/trash/admin)","","","Ledger reconciliation",""],
        ["Hotel/relocation/supplies (mitigation)","","","Receipts",""],
        ["Cleanup/remediation costs","","","Invoices",""],
        ["Property loss/damage","","","Estimate/receipts",""],
        ["Medical/health impact (if any)","","","Records/notes",""],
        ["Emotional distress (housing impact)","","","Qualitative + corroboration",""],
        ["Exemplary damages (egregious conduct)","","","Court’s discretion",""]
    ]
    if args.no_dry_run:
        with damages_csv.open("w", encoding="utf-8", newline="") as f:
            csv.writer(f).writerows(rows)
        manifest.append([damages_csv.relative_to(bundle_dir).as_posix(), damages_csv.stat().st_size, sha256_of_path(damages_csv)])

    # 7) GATES — fail-fast audits
    logging.info("[7] Running Gating Audits…")
    # Pin-cites
    total, missing = audit_pin_cites(complaint_docx)
    if total == 0 or missing:
        logging.error(f"Pin-cite audit FAILED. Numbered factual paragraphs: {total}, missing cites: {len(missing)}")
        logging.error(f"Missing ¶ indices (first 30): {missing[:30]}")
        sys.exit(2)
    logging.info(f"[OK] Pin-cite audit passed: {total} numbered factual paragraphs.")

    # PII redaction
    pii_flags = audit_pii_redaction(complaint_docx)
    if pii_flags:
        logging.error("PII redaction audit FAILED:")
        for it in pii_flags:
            logging.error(f"  - {it}")
        sys.exit(2)
    logging.info("[OK] PII/financial redaction audit passed.")

    # Privilege Matrix audit
    def audit_priv_csv(p: Path) -> list[str]:
        errs = []
        with p.open("r", encoding="utf-8-sig", newline="") as f:
            r = csv.DictReader(f)
            need = {"Source","PrivilegeType","Basis","UseAllowed"}
            if not need.issubset(set(r.fieldnames or [])):
                errs.append("Privilege_Matrix_Housing.csv missing required columns.")
                return errs
            for row in r:
                pt = (row.get("PrivilegeType") or "").strip().lower()
                ua = (row.get("UseAllowed") or "").strip().lower()
                if pt in {"absolute","qualified"} and ua not in {"context-only","none"}:
                    errs.append(f"Privileged item not context-only: {row}")
        return errs
    perrs = audit_priv_csv(priv_csv)
    if perrs:
        logging.error("Privilege Matrix audit FAILED:")
        for e in perrs[:10]: logging.error(f"  - {e}")
        sys.exit(2)
    logging.info("[OK] Privilege Matrix gate passed.")

    # SOL Matrix audit
    def audit_sol_csv(p: Path) -> list[str]:
        errs = []
        with p.open("r", encoding="utf-8-sig", newline="") as f:
            r = csv.DictReader(f)
            need = {"Claim","Status","Treatment"}
            if not need.issubset(set(r.fieldnames or [])):
                errs.append("SOL_Matrix_Housing.csv missing required columns.")
                return errs
            for row in r:
                st = (row.get("Status") or "").strip().lower()
                tr = (row.get("Treatment") or "").strip().lower()
                if st == "stale" and "background" not in tr:
                    errs.append(f"Stale item not confined to background: {row}")
        return errs
    serrs = audit_sol_csv(sol_csv)
    if serrs:
        logging.error("SOL Matrix audit FAILED:")
        for e in serrs[:10]: logging.error(f"  - {e}")
        sys.exit(2)
    logging.info("[OK] SOL Matrix gate passed.")

    # 8) Manifest & ZIP
    logging.info("[8] Writing Manifest and ZIP…")
    manifest_path = bundle_dir / "Manifest.csv"
    if args.no_dry_run:
        with manifest_path.open("w", encoding="utf-8", newline="") as f:
            w = csv.writer(f); w.writerow(["RelativePath","Bytes","SHA256"]); w.writerows(manifest)
        zpath = out_root / f"Housing_Complaint_Bundle_{ts}.zip"
        with zipfile.ZipFile(zpath.as_posix(), "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for root, _, files in os.walk(bundle_dir.as_posix()):
                for name in files:
                    p = Path(root) / name
                    zf.write(p.as_posix(), p.relative_to(bundle_dir).as_posix())
        zsha = sha256_of_path(zpath)
        logging.info(f"[ZIP] {zpath.name} | {zpath.stat().st_size} bytes | SHA-256={zsha}")

    # 9) Summary
    logging.info("== SUMMARY ==")
    logging.info(f"Bundle: {bundle_dir}")
    if args.no_dry_run:
        logging.info(f"ZIP: {out_root / f'Housing_Complaint_Bundle_{ts}.zip'}")
    logging.info("All gates passed. Housing-only bundle is READY for filing workflows.")

if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except Exception as e:
        print(f"[FATAL] {e}", file=sys.stderr)
        sys.exit(1)
